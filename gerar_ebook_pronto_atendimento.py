import os
import random
import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import geradorBancosDeQuestoesPorTopico as gb
import mysql.connector
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CATEGORIAS = [
    {
        "nome": "Clínica Médica",
        "ids_raiz": [100],
        "tem_imagem": None
    },
    {
        "nome": "Cirurgia Geral",
        "ids_raiz": [33],
        "tem_imagem": None
    },
    {
        "nome": "Pediatria",
        "ids_raiz": [48],
        "tem_imagem": None
    },
    {
        "nome": "Ortopedia",
        "ids_raiz": [261],
        "tem_imagem": None
    },
    {
        "nome": "Radiologia",
        "ids_raiz": [1121, 1505, 1618, 1619, 2308, 3421, 4018, 4052, 4630, 4692, 4811, 5928, 5964, 6035, 7013, 9603, 9666, 10008, 10285, 10426, 11533, 11692, 11743, 11998],
        "tem_imagem": 1
    },
    {
        "nome": "Psiquiatria Infantil (TEA)",
        "ids_raiz": [828],
        "tem_imagem": None
    }
]

def obter_descendentes(topicos_dict, root_ids):
    descendentes = set(root_ids)
    def traverse(node_id):
        if node_id in topicos_dict:
            for filho_id in topicos_dict[node_id]['filhos']:
                descendentes.add(filho_id)
                traverse(filho_id)
    for root_id in root_ids:
        traverse(root_id)
    return list(descendentes)

CACHE_FILE = 'cache_pronto_atendimento.json'

def carregar_cache():
    if os.path.exists(CACHE_FILE):
        try:
            with open(CACHE_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            pass
    return {"sim": [], "nao": []}

def salvar_cache(cache):
    with open(CACHE_FILE, 'w', encoding='utf-8') as f:
        json.dump(cache, f, ensure_ascii=False, indent=2)

def check_pronto_atendimento(questao, cache):
    q_id = questao['questao_id']
    if q_id in cache.get('sim', []):
        return True
    if q_id in cache.get('nao', []):
        return False
        
    texto = gb.extrair_texto_sem_imagens(questao.get('enunciado', ''))
    for alt in ['A', 'B', 'C', 'D', 'E']:
        k = f'alternativa{alt}'
        if questao.get(k):
            texto += '\n' + alt + ') ' + gb.extrair_texto_sem_imagens(questao[k])
    
    prompt = f"A seguinte questão médica aborda o tema de 'Pronto Atendimento', 'Urgência' ou 'Emergência'? Responda APENAS 'SIM' ou 'NÃO'.\n\nQuestão:\n{texto}"
    
    resposta = gb.deepseek_chat([{"role": "user", "content": prompt}], max_tokens=10)
    
    is_sim = False
    if resposta and 'SIM' in resposta.upper():
        is_sim = True
        
    if is_sim:
        cache.setdefault('sim', []).append(q_id)
    else:
        cache.setdefault('nao', []).append(q_id)
        
    salvar_cache(cache)
    return is_sim

def buscar_questoes(conn, topicos_dict, categoria, limite_busca=100):
    ids_desc = obter_descendentes(topicos_dict, categoria['ids_raiz'])
    if not ids_desc:
        return []
    
    placeholders = ','.join(['%s'] * len(ids_desc))
    query = f"""
        SELECT DISTINCT q.questao_id, q.codigo, q.enunciado, q.alternativaA, q.alternativaB, 
               q.alternativaC, q.alternativaD, q.alternativaE, q.gabarito, 
               q.comentario, q.ano, q.tem_imagem, q.instituicao, q.prova
        FROM questaoresidencia q
        JOIN classificacao_questao cq ON q.questao_id = cq.id_questao
        WHERE cq.id_topico IN ({placeholders})
          AND q.ano >= 2018
          AND q.comentario IS NOT NULL AND q.comentario != ''
          AND q.alternativaA IS NOT NULL AND q.alternativaA != ''
          AND q.alternativaB IS NOT NULL AND q.alternativaB != ''
          AND q.alternativaC IS NOT NULL AND q.alternativaC != ''
          AND q.alternativaD IS NOT NULL AND q.alternativaD != ''
    """
    
    params = list(ids_desc)
    
    if categoria['tem_imagem'] is not None:
        query += " AND q.tem_imagem = %s"
        params.append(categoria['tem_imagem'])
        
    query += " ORDER BY RAND() LIMIT %s"
    params.append(limite_busca)
    
    cursor = conn.cursor(dictionary=True)
    cursor.execute(query, tuple(params))
    questoes = cursor.fetchall()
    cursor.close()
    
    return questoes

def add_heading_with_style(doc, text, level):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0, 51, 102)
    return heading

def limpar_html_para_docx(html_text):
    if not html_text:
        return ""
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html_text, 'html.parser')
    return soup.get_text(separator='\n').strip()

def add_questao_to_doc(doc, q_num, questao):
    enunciado = limpar_html_para_docx(questao.get('enunciado', ''))
    p = doc.add_paragraph()
    run = p.add_run(f"Questão {q_num} - {questao.get('instituicao', 'Instituição')} ({questao.get('ano', '')})\n")
    run.bold = True
    
    p.add_run(enunciado + "\n")
    
    # Tentativa simples de inserir imagens do enunciado se houver
    if questao.get('tem_imagem') == 1:
        # Nota: gb.add_imagens_enunciado geralmente requer o ID da questão e a pasta de imagens
        # Vou chamar a função da lib se existir
        pass # Por simplificação, se tiver imagem online ou no texto HTML, gb.add_imagens_enunciado cuida disso
        # Vamos usar a lógica de adicionar imagens
        # Requer baixar imagens.
    
    for alt in ['A', 'B', 'C', 'D', 'E']:
        k = f'alternativa{alt}'
        if questao.get(k):
            alt_text = limpar_html_para_docx(questao[k])
            p.add_run(f"{alt}) {alt_text}\n")
    
    p.add_run("\n")

def gerar_ebook():
    total_desejado = 20
    try:
        entrada = input("Informe a quantidade total de questões desejada para o E-book (padrão 20): ").strip()
        if entrada:
            total_desejado = int(entrada)
    except ValueError:
        print("[AVISO] Valor inválido. Utilizando o padrão 20.")
        
    num_categorias = len(CATEGORIAS)
    qtd_base = max(1, total_desejado // num_categorias)
    resto = total_desejado % num_categorias

    conn = gb.get_connection()
    topicos_dict, topicos_raiz = gb.carregar_hierarquia_topicos(conn)
    
    cache = carregar_cache()
    questoes_finais = []
    
    for i, cat in enumerate(CATEGORIAS):
        qtd_alvo = qtd_base + (1 if i < resto else 0)
        
        print(f"\n[LOG] Buscando questões para {cat['nome']} (Alvo: {qtd_alvo})...")
        questoes_cat = buscar_questoes(conn, topicos_dict, cat, limite_busca=max(100, qtd_alvo * 15))
        print(f"[LOG] {len(questoes_cat)} questões encontradas no BD para {cat['nome']}. Filtrando por PA...")
        
        selecionadas = []
        for q in questoes_cat:
            if check_pronto_atendimento(q, cache):
                print(f"  -> Questão {q['questao_id']} é Pronto Atendimento!")
                selecionadas.append(q)
                if len(selecionadas) == qtd_alvo:
                    break
        
        print(f"[LOG] Total selecionadas para {cat['nome']}: {len(selecionadas)}")
        
        # Adicionar as selecionadas
        for q in selecionadas:
            q['_categoria'] = cat['nome']
        questoes_finais.extend(selecionadas)
    
    conn.close()
    
    # Criar documento DOCX
    doc = Document()
    gb.configurar_metadados_documento(doc, len(questoes_finais))
    
    add_heading_with_style(doc, 'E-book: Casos de Pronto Atendimento', 0)
    doc.add_paragraph("Este material contém questões selecionadas com foco em Pronto Atendimento, Urgência e Emergência, agrupadas por grandes áreas.")
    doc.add_page_break()
    
    q_num = 1
    gabaritos = []
    
    for cat in CATEGORIAS:
        cat_questoes = [q for q in questoes_finais if q['_categoria'] == cat['nome']]
        if not cat_questoes:
            continue
            
        add_heading_with_style(doc, cat['nome'], 1)
        for q in cat_questoes:
            # Para inserir imagens, vamos usar gb.add_imagens_enunciado
            
            p = doc.add_paragraph()
            run = p.add_run(f"Questão {q_num} - {q.get('instituicao', '')} ({q.get('ano', '')})\n")
            run.bold = True
            
            # Enunciado pode ter HTML. No script original, eles inserem via bs4 + formatacao
            enunciado_limpo = limpar_html_para_docx(q.get('enunciado', ''))
            p.add_run(enunciado_limpo + "\n")
            
            img_dir = r'C:\Users\elman\GoogleDrive\QuestoesMED\imagens\QuestoesResidencia'
            codigo_q = q.get('codigo') or str(q['questao_id'])
            gb.add_imagens_enunciado(doc, q.get('enunciado', ''), codigo_q, img_dir)
            
            # Alternativas
            p_alt = doc.add_paragraph()
            gabarito_correto = q.get('gabarito', '').upper()
            for alt in ['A', 'B', 'C', 'D', 'E']:
                k = f'alternativa{alt}'
                if q.get(k):
                    alt_text = limpar_html_para_docx(q[k])
                    run = p_alt.add_run(f"{alt}) {alt_text}\n")
                    if alt == gabarito_correto:
                        run.bold = True
            
            gabaritos.append({
                'num': q_num,
                'gabarito': gabarito_correto,
                'comentario': q.get('comentario', ''),
                'id': q['questao_id'],
                'codigo': q.get('codigo'),
                'cat': cat['nome']
            })
            q_num += 1
            
            doc.add_paragraph("\n")
            
    doc.add_page_break()
    add_heading_with_style(doc, "Gabarito e Comentários", 1)
    
    for g in gabaritos:
        p = doc.add_paragraph()
        run = p.add_run(f"Questão {g['num']} ({g['cat']}) - Gabarito: {g['gabarito']}\n")
        run.bold = True
        
        coment = limpar_html_para_docx(g['comentario']) if g['comentario'] else "Sem comentário."
        p.add_run(coment + "\n")
        
        img_dir = r'C:\Users\elman\GoogleDrive\QuestoesMED\imagens\QuestoesResidencia_comentarios'
        codigo_g = g.get('codigo') or str(g['id'])
        gb.add_imagens_enunciado(doc, g['comentario'] or '', codigo_g, img_dir)
        
    out_path = os.path.join(os.path.dirname(__file__), 'Ebook_Pronto_Atendimento.docx')
    doc.save(out_path)
    print(f"\n[SUCESSO] E-book gerado em: {out_path}")

if __name__ == '__main__':
    gerar_ebook()
