import sys
import mysql.connector
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION
from markdown2 import markdown
from bs4 import BeautifulSoup
from bs4 import Comment
import os
import re
from docx.shared import Inches
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.image.exceptions import UnrecognizedImageError
import mimetypes
from datetime import datetime
from PIL import Image
import requests
import json

# Configurações do banco
DB_CONFIG = {
    "host": "localhost",
    "user": "root",
    "password": "El@mysql.32",
    "database": "qconcursos"
}

# Configurações da API DeepSeek
def load_api_key():
    """
    Carrega a API key do arquivo api_key.txt.
    """
    api_key_path = os.path.join(os.path.dirname(__file__), 'api_key.txt')
    try:
        with open(api_key_path, 'r', encoding='utf-8') as f:
            api_key = f.read().strip()
            if not api_key:
                raise ValueError("API key está vazia no arquivo api_key.txt")
            return api_key
    except FileNotFoundError:
        print(f"[ERRO] Arquivo api_key.txt não encontrado em {api_key_path}")
        print("[ERRO] Crie o arquivo api_key.txt na raiz do projeto com sua API key do DeepSeek")
        raise
    except Exception as e:
        print(f"[ERRO] Erro ao ler api_key.txt: {str(e)}")
        raise

DEEPSEEK_CONFIG = {
    "api_key": load_api_key(),
    "model": "deepseek-chat",
    "temperature": 0.1,
    "url": "https://api.deepseek.com/v1/chat/completions"
}

def deepseek_chat(messages, max_tokens=200, temperature=None):
    """
    Executa uma chamada à API DeepSeek com os parâmetros padrão do projeto.
    Retorna o conteúdo textual da resposta da IA ou None em caso de erro.
    """
    if temperature is None:
        temperature = DEEPSEEK_CONFIG["temperature"]

    headers = {
        "Authorization": f"Bearer {DEEPSEEK_CONFIG['api_key']}",
        "Content-Type": "application/json"
    }

    payload = {
        "model": DEEPSEEK_CONFIG["model"],
        "messages": messages,
        "temperature": temperature,
        "max_tokens": max_tokens
    }

    try:
        response = requests.post(DEEPSEEK_CONFIG["url"], headers=headers, json=payload)
        response.raise_for_status()
        data = response.json()
        content = data['choices'][0]['message']['content']
        if content is None:
            print("[ERRO] Resposta da API DeepSeek sem conteúdo.")
            return None
        return content
    except requests.exceptions.RequestException as e:
        print(f"[ERRO] Erro na chamada da API DeepSeek: {str(e)}")
        return None
    except (KeyError, IndexError) as e:
        print(f"[ERRO] Estrutura inesperada na resposta da API DeepSeek: {str(e)}")
        return None
    except Exception as e:
        print(f"[ERRO] Falha inesperada ao processar resposta da API DeepSeek: {str(e)}")
        return None

def verificar_e_adicionar_imagem(document, img_path, max_width=None):
    """
    Função auxiliar para verificar e adicionar imagem de forma segura.
    Preserva o tamanho original se for menor que max_width.
    Retorna True se a imagem foi adicionada com sucesso, False caso contrário.
    """
    try:
        # Verificar se o arquivo existe
        if not os.path.exists(img_path):
            print(f"[AVISO] Arquivo de imagem não encontrado: {img_path}")
            return False
        
        # Verificar se é um arquivo válido
        if not os.path.isfile(img_path):
            print(f"[AVISO] Caminho não é um arquivo válido: {img_path}")
            return False
        
        # Verificar tamanho do arquivo
        file_size = os.path.getsize(img_path)
        if file_size == 0:
            print(f"[AVISO] Arquivo de imagem vazio: {img_path}")
            return False
        
        # Verificar formato da imagem
        mime_type, _ = mimetypes.guess_type(img_path)
        if mime_type and not mime_type.startswith('image/'):
            print(f"[AVISO] Arquivo não parece ser uma imagem válida: {img_path} (tipo: {mime_type})")
            return False
        
        # Calcular tamanho adequado da imagem
        width_to_use = None
        if max_width:
            try:
                with Image.open(img_path) as img:
                    original_width_px = img.width
                    original_height_px = img.height
                    
                    # Converter largura original de pixels para inches (assumindo 96 DPI)
                    original_width_inches = Inches(original_width_px / 96.0)
                    
                    # Usar o menor valor entre largura original e largura máxima
                    width_to_use = min(original_width_inches, max_width)
                    
                    status = "ORIGINAL" if width_to_use == original_width_inches else "REDUZIDA"
                    print(f"[LOG] Imagem {os.path.basename(img_path)}: {original_width_px}x{original_height_px}px ({status})")
                    print(f"[LOG] Largura: original={original_width_inches:.2f}in, máx={max_width:.2f}in, usada={width_to_use:.2f}in")
                    
            except Exception as e:
                print(f"[AVISO] Não foi possível obter dimensões da imagem {img_path}: {str(e)}")
                width_to_use = max_width  # Fallback para largura máxima
        
        # Tentar adicionar a imagem
        if width_to_use:
            document.add_picture(img_path, width=width_to_use)
        else:
            document.add_picture(img_path)
        
        print(f"[LOG] Imagem adicionada com sucesso: {img_path}")
        return True
        
    except UnrecognizedImageError as e:
        print(f"[ERRO] Formato de imagem não reconhecido: {img_path}")
        print(f"[ERRO] Detalhes: {str(e)}")
        return False
    except Exception as e:
        print(f"[ERRO] Erro ao adicionar imagem {img_path}: {str(e)}")
        return False

def get_connection():
    print("[LOG] Abrindo conexão com o banco de dados...")
    return mysql.connector.connect(**DB_CONFIG)

def identificar_questoes_incompletas(conn, instituicao, resto_mod5=0):
    """
    Identifica questões com comentários incompletos que terminam com 'analisar as alternativas' ou 'analisar as opções',
    que ainda não possuem resposta gerada pela IA (gabaritoIA IS NULL) e pertencem à instituição informada.
    Retorna lista de questões do conjunto [INCOMPLETO].
    """
    print("[LOG] Identificando questões com comentários incompletos...")
    
    cursor = conn.cursor(dictionary=True)
    
    # Buscar questões que não foram respondidas pela IA e sejam de determinada instituição
    query = """
    SELECT questao_id, codigo, enunciado, alternativaA, alternativaB, alternativaC,
           alternativaD, alternativaE, gabarito, comentario
    FROM questaoresidencia
    WHERE gabaritoIA IS NULL
      AND comentarioIA IS NULL
      AND instituicao LIKE %s
      AND (MOD(questao_id, 5) = %s)
    ORDER BY questao_id
    """

    instituicao_like = f"%{instituicao}%"
    cursor.execute(query, (instituicao_like, resto_mod5))
    questoes_incompletas = cursor.fetchall()
    
    print(f"[LOG] Encontradas {len(questoes_incompletas)} questões com comentários incompletos")
    
    return questoes_incompletas

def chamar_api_deepseek(enunciado, alternativas, gabarito_correto):
    """
    Chama a API DeepSeek para analisar uma questão e obter resposta e justificativa.
    """
    print(f"[LOG] Chamando API DeepSeek para questão...")
    
    # Montar o texto da questão
    texto_questao = f"Enunciado: {enunciado}\n\n"
    for alt in ['A', 'B', 'C', 'D', 'E']:
        if alternativas.get(f'alternativa{alt}'):
            texto_questao += f"{alt}) {alternativas[f'alternativa{alt}']}\n"
    
    # Primeira chamada: solicitar apenas a resposta
    prompt_resposta = f"""
Analise a seguinte questão de medicina e responda APENAS com a letra da alternativa correta (A, B, C, D ou E).

{texto_questao}

Responda apenas com a letra da alternativa correta:
"""
    
    try:
        resposta_bruta = deepseek_chat(
            [{"role": "user", "content": prompt_resposta}],
            max_tokens=10
        )
        if not resposta_bruta:
            print("[AVISO] Nenhuma resposta recebida da IA para a alternativa correta.")
            return None, None, None

        resposta_ia = resposta_bruta.strip().upper()
        
        # Verificar se a resposta é válida (A, B, C, D ou E)
        if resposta_ia not in ['A', 'B', 'C', 'D', 'E']:
            print(f"[AVISO] Resposta inválida da IA: {resposta_ia}")
            return None, None, None
        
        print(f"[LOG] IA respondeu: {resposta_ia}, Gabarito correto: {gabarito_correto}")
        
        # Verificar se acertou
        if resposta_ia == gabarito_correto.upper():
            print(f"[LOG] IA acertou! Solicitando justificativa...")
            
            # Segunda chamada: solicitar justificativa detalhada
            prompt_justificativa = f"""
A questão anterior foi respondida corretamente. Agora forneça uma justificativa detalhada e educativa em formato markdown, incluindo:

1. Explicação clara do conceito médico envolvido
2. Análise de cada alternativa (por que está correta ou incorreta)
3. Uso de recursos visuais como tabelas, emojis e formatação markdown
4. Estrutura organizada com títulos e seções

{texto_questao}

Resposta correta: {resposta_ia}

Forneça a justificativa completa em markdown:
"""
            
            justificativa = deepseek_chat(
                [{"role": "user", "content": prompt_justificativa}],
                max_tokens=2000
            )
            if justificativa:
                justificativa = justificativa.strip()
            else:
                print("[AVISO] Não foi possível obter justificativa da IA.")
            
            return resposta_ia, justificativa, True  # acertou = True
        else:
            print(f"[LOG] IA errou. Resposta: {resposta_ia}, Gabarito: {gabarito_correto}")
            return resposta_ia, None, False  # acertou = False
            
    except Exception as e:
        print(f"[ERRO] Erro inesperado na API: {str(e)}")
        return None, None, None

def extrair_primeiro_inteiro(texto):
    """
    Extrai o primeiro número inteiro encontrado na string fornecida.
    """
    if not texto:
        return None
    match = re.search(r'\d+', texto)
    if not match:
        return None
    try:
        return int(match.group())
    except ValueError:
        return None

def carregar_hierarquia_topicos(conn):
    """
    Carrega toda a hierarquia de tópicos e organiza filhos para navegação.
    Retorna um dicionário {id_topico: {...}} e uma lista ordenada de tópicos raiz.
    """
    print("[LOG] Carregando hierarquia completa de tópicos...")
    cursor = conn.cursor(dictionary=True)
    cursor.execute("SELECT id, nome, id_pai FROM topico")
    topicos = cursor.fetchall()
    cursor.close()

    if not topicos:
        print("[ERRO] Nenhum tópico encontrado na base de dados.")
        return {}, []

    topicos_dict = {}
    for t in topicos:
        topicos_dict[t['id']] = {
            'id': t['id'],
            'nome': t['nome'],
            'id_pai': t['id_pai'],
            'filhos': []
        }

    for t in topicos:
        id_pai = t['id_pai']
        if id_pai and id_pai in topicos_dict:
            topicos_dict[id_pai]['filhos'].append(t['id'])

    for info in topicos_dict.values():
        info['filhos'] = sorted(
            [fid for fid in info['filhos'] if fid in topicos_dict],
            key=lambda tid: topicos_dict[tid]['nome']
        )

    topicos_raiz = sorted(
        [t_id for t_id, t_info in topicos_dict.items() if t_info['id_pai'] is None],
        key=lambda tid: topicos_dict[tid]['nome']
    )

    print(f"[LOG] Hierarquia carregada: {len(topicos_dict)} tópicos, {len(topicos_raiz)} raízes.")
    return topicos_dict, topicos_raiz

def buscar_questoes_sem_classificacao(conn, limite=None, filtro_instituicao=None, resto_mod5=None, filtro_ano=None):
    """
    Recupera questões que ainda não possuem registros em classificacao_questao.
    Permite filtrar por instituição e por resto de divisão por 5.
    """
    print("[LOG] Buscando questões sem classificação...")
    cursor = conn.cursor(dictionary=True)
    query = """
    SELECT q.questao_id, q.codigo, q.enunciado, q.gabarito, COALESCE(q.gabarito_texto, '') AS gabarito_texto,
           COALESCE(q.instituicao, '') AS instituicao
    FROM questaoresidencia q
    WHERE NOT EXISTS (
        SELECT 1
        FROM classificacao_questao cq
        WHERE cq.id_questao = q.questao_id
    )
    """
    params = []

    if filtro_instituicao:
        query += " AND q.instituicao LIKE %s"
        params.append(f"%{filtro_instituicao}%")

    if resto_mod5 is not None:
        query += " AND MOD(q.questao_id, 5) = %s"
        params.append(resto_mod5)

    if filtro_ano is not None:
        query += " AND q.ano >= %s"
        params.append(filtro_ano)

    query += " ORDER BY q.questao_id"

    if limite and limite > 0:
        query += " LIMIT %s"
        params.append(limite)

    cursor.execute(query, tuple(params))
    questoes = cursor.fetchall()
    cursor.close()

    print(f"[LOG] Questões encontradas sem classificação: {len(questoes)}")
    return questoes

def montar_lista_opcoes(topicos_dict, topico_ids):
    """
    Monta a lista textual das opções e o mapeamento índice->id_topico.
    """
    opcoes_texto = []
    mapa_indice = {}
    for idx, topico_id in enumerate(topico_ids, start=1):
        nome = topicos_dict[topico_id]['nome']
        opcoes_texto.append(f"{idx}. {nome}")
        mapa_indice[idx] = topico_id
    return "\n".join(opcoes_texto), mapa_indice

def classificar_questao_hierarquica(questao, topicos_dict, topicos_raiz_ids):
    """
    Realiza a classificação hierárquica de uma questão utilizando a API DeepSeek.
    Retorna (lista_de_topicos, classificacao_completa) onde classificacao_completa indica
    se foi possível chegar até um tópico folha.
    """
    enunciado_limpo = extrair_texto_sem_imagens(questao.get('enunciado', '') or '')
    gabarito_texto = questao.get('gabarito_texto') or ''
    gabarito_texto_limpo = extrair_texto_sem_imagens(gabarito_texto) if gabarito_texto else ''

    enunciado_limpo = enunciado_limpo[:4000]
    gabarito_composto = (questao.get('gabarito') or '').strip()
    if gabarito_texto_limpo:
        gabarito_composto = f"{gabarito_composto} - {gabarito_texto_limpo}"
    gabarito_composto = gabarito_composto[:1000]

    caminho_topicos = []
    opcoes_atual = list(topicos_raiz_ids)
    nivel = 1
    visitados = set()
    classificacao_completa = True

    while opcoes_atual:
        lista_opcoes, mapa_indice = montar_lista_opcoes(topicos_dict, opcoes_atual)
        if not mapa_indice:
            print("[AVISO] Lista de tópicos vazia durante a classificação.")
            classificacao_completa = False
            break

        prompt_base = (
            "Classifique a questão abaixo em um dos tópicos listados a seguir. "
            "Informe APENAS o número correspondente ao tópico que melhor define o assunto da questão.\n\n"
            f"[ENUNCIADO]: {enunciado_limpo}\n\n"
            f"[GABARITO]: {gabarito_composto}\n\n"
            f"[TÓPICOS NÍVEL {nivel}]:\n{lista_opcoes}\n"
        )

        numero_escolhido = None
        for tentativa in range(2):
            prompt = prompt_base if tentativa == 0 else (
                prompt_base +
                f"\nResponda apenas com um número da lista acima. Opções válidas: {', '.join(str(i) for i in mapa_indice.keys())}."
            )

            resposta = deepseek_chat(
                [{"role": "user", "content": prompt}],
                max_tokens=10
            )

            if not resposta:
                print("[AVISO] Sem resposta da IA para seleção de tópico.")
                classificacao_completa = False
                numero_escolhido = None
                break

            numero = extrair_primeiro_inteiro(resposta)
            if numero in mapa_indice:
                numero_escolhido = numero
                break

            print(f"[AVISO] Resposta inválida para seleção de tópico: '{resposta}'. Tentativa {tentativa + 1}/2.")

        if numero_escolhido is None:
            if caminho_topicos:
                print("[AVISO] Não foi possível avançar para um nível mais específico. Usando classificação parcial.")
                classificacao_completa = False
            else:
                print("[ERRO] Não foi possível obter um número válido para o tópico inicial.")
            break

        topico_id = mapa_indice[numero_escolhido]
        nome_topico = topicos_dict[topico_id]['nome']
        print(f"[LOG] Nível {nivel}: selecionado tópico {numero_escolhido} -> ID {topico_id} ({nome_topico})")

        if topico_id in visitados:
            print(f"[AVISO] Ciclo detectado na hierarquia de tópicos (ID {topico_id}). Classificação parcial será utilizada.")
            classificacao_completa = False
            break

        caminho_topicos.append(topico_id)
        visitados.add(topico_id)

        filhos = topicos_dict[topico_id]['filhos']
        if not filhos:
            print(f"[LOG] Tópico {topico_id} não possui filhos. Classificação encerrada.")
            break

        opcoes_atual = filhos
        nivel += 1

    if not caminho_topicos:
        print("[AVISO] Nenhum tópico foi selecionado para a questão.")
        return [], False

    return caminho_topicos, classificacao_completa

def inserir_classificacao_questao(conn, questao_id, topico_ids):
    """
    Insere registros na tabela classificacao_questao para os tópicos fornecidos.
    Usa INSERT IGNORE para evitar duplicidades.
    """
    cursor = conn.cursor()
    try:
        for topico_id in topico_ids:
            cursor.execute(
                "INSERT IGNORE INTO classificacao_questao (id_questao, id_topico) VALUES (%s, %s)",
                (questao_id, topico_id)
            )
        conn.commit()
    except Exception as e:
        conn.rollback()
        raise e
    finally:
        cursor.close()

def processar_classificacao_questoes_sem_topico(conn, limite=None, filtro_instituicao=None, resto_mod5=None, filtro_ano=None):
    """
    Processa questões sem classificação hierárquica utilizando a API DeepSeek.
    """
    topicos_dict, topicos_raiz = carregar_hierarquia_topicos(conn)
    if not topicos_raiz:
        print("[ERRO] Não foi possível carregar tópicos raiz. Processo interrompido.")
        return

    questoes = buscar_questoes_sem_classificacao(
        conn,
        limite=limite,
        filtro_instituicao=filtro_instituicao,
        resto_mod5=resto_mod5,
        filtro_ano=filtro_ano
    )

    if not questoes:
        print("[LOG] Nenhuma questão sem classificação encontrada para os filtros informados.")
        return

    sucessos = 0
    falhas = 0
    classificacoes_completas = 0
    classificacoes_parciais = 0

    for idx, questao in enumerate(questoes, start=1):
        print("\n" + "=" * 80)
        print(f"[LOG] Classificando questão {idx}/{len(questoes)} | ID {questao['questao_id']} | Código {questao.get('codigo', '')}")
        if questao.get('instituicao'):
            print(f"[LOG] Instituição: {questao['instituicao']}")

        caminho_topicos, classificacao_completa = classificar_questao_hierarquica(questao, topicos_dict, topicos_raiz)

        if not caminho_topicos:
            print(f"[ERRO] Falha na classificação da questão {questao['questao_id']}.")
            falhas += 1
            continue

        nomes_topicos = " > ".join(topicos_dict[t]['nome'] for t in caminho_topicos)
        print(f"[LOG] Caminho de tópicos selecionado: {nomes_topicos}")
        if not classificacao_completa:
            print("[INFO] Classificação parcial registrada (tópico folha não identificado).")

        try:
            inserir_classificacao_questao(conn, questao['questao_id'], caminho_topicos)
            sucessos += 1
            if classificacao_completa:
                classificacoes_completas += 1
            else:
                classificacoes_parciais += 1
            print(f"[SUCESSO] Classificação salva para a questão {questao['questao_id']}.")
        except Exception as e:
            print(f"[ERRO] Falha ao salvar classificação da questão {questao['questao_id']}: {str(e)}")
            falhas += 1

    print("\n[LOG] === RESUMO DO MODO 6 ===")
    print(f"[LOG] Questões processadas: {len(questoes)}")
    print(f"[LOG] Classificações salvas: {sucessos}")
    print(f"[LOG]   - Completas: {classificacoes_completas}")
    print(f"[LOG]   - Parciais: {classificacoes_parciais}")
    print(f"[LOG] Falhas: {falhas}")

def processar_questoes_incompletas(conn, instituicao, resto_mod5=0):
    """
    Processa todas as questões incompletas de uma instituição específica usando a API DeepSeek.
    """
    print("[LOG] === MODO 4: Processando questões com comentários incompletos ===")
    
    # Identificar questões incompletas
    questoes_incompletas = identificar_questoes_incompletas(conn, instituicao, resto_mod5)
    
    if not questoes_incompletas:
        print("[LOG] Nenhuma questão incompleta encontrada.")
        return
    
    print(f"[LOG] Processando {len(questoes_incompletas)} questões incompletas...")
    
    cursor = conn.cursor()
    sucessos = 0
    erros = 0
    acertos_ia = 0
    erros_ia = 0
    questoes_avaliadas = 0
    
    for i, questao in enumerate(questoes_incompletas, 1):
        print(f"\n[LOG] Processando questão {i}/{len(questoes_incompletas)}: {questao['codigo']}")
        
        # Preparar alternativas
        alternativas = {
            'alternativaA': questao.get('alternativaA', ''),
            'alternativaB': questao.get('alternativaB', ''),
            'alternativaC': questao.get('alternativaC', ''),
            'alternativaD': questao.get('alternativaD', ''),
            'alternativaE': questao.get('alternativaE', '')
        }
        
        # Chamar API DeepSeek
        resposta_ia, justificativa, acertou = chamar_api_deepseek(
            questao['enunciado'], 
            alternativas, 
            questao['gabarito']
        )
        
        if resposta_ia is None:
            print(f"[ERRO] Falha na análise da questão {questao['codigo']}")
            erros += 1
            continue
        
        questoes_avaliadas += 1
        if acertou:
            acertos_ia += 1
        else:
            erros_ia += 1
        
        # Preparar dados para atualização
        data_atual = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        autor = "DeepSeek AI"
        
        try:
            if acertou and justificativa:
                # IA acertou - atualizar com justificativa completa
                update_query = """
                UPDATE questaoresidencia 
                SET comentarioIA = %s, 
                    comentario_autor = %s, 
                    comentario_data = %s, 
                    gabaritoIA = %s
                WHERE questao_id = %s
                """
                cursor.execute(update_query, (
                    justificativa, 
                    autor, 
                    data_atual, 
                    resposta_ia, 
                    questao['questao_id']
                ))
                print(f"[SUCESSO] Questão {questao['codigo']} atualizada com justificativa completa")
            else:
                # IA errou - atualizar apenas com dados básicos
                update_query = """
                UPDATE questaoresidencia 
                SET comentario_autor = %s, 
                    comentario_data = %s, 
                    gabaritoIA = %s
                WHERE questao_id = %s
                """
                cursor.execute(update_query, (
                    autor, 
                    data_atual, 
                    resposta_ia, 
                    questao['questao_id']
                ))
                print(f"[INFO] Questão {questao['codigo']} atualizada (IA errou)")
            
            # Commit após cada questão atualizada
            conn.commit()
            sucessos += 1
            
        except Exception as e:
            print(f"[ERRO] Falha ao atualizar questão {questao['codigo']}: {str(e)}")
            # Rollback apenas da operação atual
            try:
                conn.rollback()
            except Exception:
                pass
            erros += 1
    
    # Fazer commit das alterações 
    try:
        conn.commit()
        print(f"\n[LOG] === RESUMO DO MODO 4 ===")
        print(f"[LOG] Questões processadas: {len(questoes_incompletas)}")
        print(f"[LOG] Sucessos: {sucessos}")
        print(f"[LOG] Erros: {erros}")
        print(f"[LOG] Questões avaliadas pela IA: {questoes_avaliadas}")
        if questoes_avaliadas > 0:
            taxa_acerto = (acertos_ia / questoes_avaliadas) * 100
            print(f"[LOG] Taxa de acerto da IA: {taxa_acerto:.2f}% ({acertos_ia}/{questoes_avaliadas})")
        else:
            print(f"[LOG] Taxa de acerto da IA: N/A (nenhuma questão avaliada)")
        print(f"[LOG] Total de erros da IA: {erros_ia}")
        print(f"[LOG] Alterações commitadas no banco de dados")
    except Exception as e:
        print(f"[ERRO] Falha ao fazer commit: {str(e)}")
        conn.rollback()

def processar_questoes_por_id(conn, questao_ids=None, limite=None, filtro_instituicao=None, resto_mod5=None, filtro_ano=None):
    """
    Processa questões usando a API DeepSeek.
    Atualiza as colunas gabaritoIA e comentarioIA.
    
    Pode filtrar por:
    - Lista de IDs de questões (questao_ids)
    - Limite de questões (limite)
    - Instituição (filtro_instituicao)
    - Resto módulo 5 (resto_mod5)
    - Ano mínimo (filtro_ano)
    
    Se questao_ids for fornecido, será combinado com os outros filtros.
    Se questao_ids não for fornecido, usará apenas os outros filtros.
    """
    print("[LOG] === MODO 5: Processando questões ===")
    
    # Capturar horário inicial
    horario_inicial = datetime.now()
    horario_inicial_str = horario_inicial.strftime('%Y-%m-%d %H:%M:%S')
    
    # Verificar se pelo menos um critério foi fornecido
    if not questao_ids and limite is None and filtro_instituicao is None and resto_mod5 is None and filtro_ano is None:
        print("[ERRO] Nenhum critério de busca fornecido. Forneça IDs ou pelo menos um filtro.")
        return
    
    cursor = conn.cursor(dictionary=True)
    
    # Construir query dinamicamente
    query = """
    SELECT questao_id, codigo, enunciado, alternativaA, alternativaB, alternativaC, 
           alternativaD, alternativaE, gabarito, comentario
    FROM questaoresidencia 
    WHERE 1=1
    """
    params = []
    
    # Filtrar por IDs se fornecido
    if questao_ids:
        format_strings = ','.join(['%s'] * len(questao_ids))
        query += f" AND questao_id IN ({format_strings})"
        params.extend(questao_ids)
    
    # Filtrar por instituição
    if filtro_instituicao:
        query += " AND instituicao LIKE %s"
        params.append(f"%{filtro_instituicao}%")
    
    # Filtrar por resto módulo 5
    if resto_mod5 is not None:
        query += " AND MOD(questao_id, 5) = %s"
        params.append(resto_mod5)
    
    # Filtrar por ano mínimo
    if filtro_ano is not None:
        query += " AND ano >= %s"
        params.append(filtro_ano)
    
    query += " ORDER BY questao_id"
    
    # Aplicar limite se fornecido
    if limite and limite > 0:
        query += " LIMIT %s"
        params.append(limite)
    
    cursor.execute(query, tuple(params))
    questoes = cursor.fetchall()
    
    if not questoes:
        print(f"[ERRO] Nenhuma questão encontrada com os critérios fornecidos.")
        return
    
    print(f"[LOG] Encontradas {len(questoes)} questão(ões) no banco de dados")
    
    # Verificar se há IDs que não foram encontrados (apenas se IDs foram fornecidos)
    ids_nao_encontrados = []
    if questao_ids:
        ids_encontrados = {q['questao_id'] for q in questoes}
        ids_nao_encontrados = [id_q for id_q in questao_ids if id_q not in ids_encontrados]
        if ids_nao_encontrados:
            print(f"[AVISO] Os seguintes IDs não foram encontrados: {ids_nao_encontrados}")
    
    cursor.close()
    cursor = conn.cursor()
    sucessos = 0
    erros = 0
    acertos = 0
    
    for i, questao in enumerate(questoes, 1):
        print(f"\n[LOG] Processando questão {i}/{len(questoes)}: ID {questao['questao_id']} (Código: {questao['codigo']})")
        
        # Preparar alternativas
        alternativas = {
            'alternativaA': questao.get('alternativaA', ''),
            'alternativaB': questao.get('alternativaB', ''),
            'alternativaC': questao.get('alternativaC', ''),
            'alternativaD': questao.get('alternativaD', ''),
            'alternativaE': questao.get('alternativaE', '')
        }
        
        # Chamar API DeepSeek
        resposta_ia, justificativa, acertou = chamar_api_deepseek(
            questao['enunciado'], 
            alternativas, 
            questao['gabarito']
        )
        
        if resposta_ia is None:
            print(f"[ERRO] Falha na análise da questão ID {questao['questao_id']}")
            erros += 1
            continue
        
        # Preparar dados para atualização
        data_atual = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        autor = "DeepSeek AI"
        
        try:
            # Contar acertos da IA (independente de ter justificativa)
            if acertou:
                acertos += 1
            
            if acertou and justificativa:
                # IA acertou - atualizar com justificativa completa
                update_query = """
                UPDATE questaoresidencia 
                SET comentarioIA = %s, 
                    comentario_autor = %s, 
                    comentario_data = %s, 
                    gabaritoIA = %s
                WHERE questao_id = %s
                """
                cursor.execute(update_query, (
                    justificativa, 
                    autor, 
                    data_atual, 
                    resposta_ia, 
                    questao['questao_id']
                ))
                print(f"[SUCESSO] Questão ID {questao['questao_id']} (Código: {questao['codigo']}) atualizada com justificativa completa")
            else:
                # IA errou - atualizar apenas com dados básicos
                update_query = """
                UPDATE questaoresidencia 
                SET comentario_autor = %s, 
                    comentario_data = %s, 
                    gabaritoIA = %s
                WHERE questao_id = %s
                """
                cursor.execute(update_query, (
                    autor, 
                    data_atual, 
                    resposta_ia, 
                    questao['questao_id']
                ))
                print(f"[INFO] Questão ID {questao['questao_id']} (Código: {questao['codigo']}) atualizada (IA errou)")
            
            # Commit após cada questão atualizada
            conn.commit()
            sucessos += 1
            
        except Exception as e:
            print(f"[ERRO] Falha ao atualizar questão ID {questao['questao_id']}: {str(e)}")
            # Rollback apenas da operação atual
            try:
                conn.rollback()
            except Exception:
                pass
            erros += 1
    
    # Capturar horário final e calcular duração
    horario_final = datetime.now()
    horario_final_str = horario_final.strftime('%Y-%m-%d %H:%M:%S')
    duracao_total = horario_final - horario_inicial
    duracao_minutos = duracao_total.total_seconds() / 60
    
    print(f"\n[LOG] === RESUMO DO MODO 5 ===")
    print(f"[LOG] Questões processadas: {len(questoes)}")
    print(f"[LOG] Sucessos: {sucessos}")
    print(f"[LOG] Erros: {erros}")
    
    # Calcular e exibir taxa de acerto da IA
    if sucessos > 0:
        taxa_acerto = (acertos / sucessos) * 100
        print(f"[LOG] Acertos da IA: {acertos}/{sucessos} ({taxa_acerto:.2f}%)")
    else:
        print(f"[LOG] Acertos da IA: 0 (nenhuma questão processada com sucesso)")
    
    if ids_nao_encontrados:
        print(f"[LOG] IDs não encontrados: {len(ids_nao_encontrados)}")
    
    # Exibir informações de tempo
    print(f"[LOG] Horário inicial: {horario_inicial_str}")
    print(f"[LOG] Horário final: {horario_final_str}")
    print(f"[LOG] Duração da execução: {duracao_minutos:.2f} minutos")

def get_topic_tree_recursive(conn, id_topico, current_level=1, max_level=4):
    print(f"[LOG] Buscando árvore de tópicos recursivamente para id_topico={id_topico} (nível {current_level})")
    cursor = conn.cursor(dictionary=True)
    cursor.execute("SELECT id, nome FROM topico WHERE id = %s", (id_topico,))
    root = cursor.fetchone()
    
    if not root:
        return None
    
    # Adicionar campo 'children' e 'nivel' se não existir
    root['children'] = []
    root['nivel'] = current_level
    
    if current_level >= max_level:
        print(f"[LOG] Limite de profundidade atingido (nível {current_level}) para tópico {root['nome']}")
        return root
    
    cursor.execute("SELECT id, nome FROM topico WHERE id_pai = %s", (id_topico,))
    children = cursor.fetchall()
    for child in children:
        child_tree = get_topic_tree_recursive(conn, child['id'], current_level + 1, max_level)
        if child_tree:
            root['children'].append(child_tree)
    
    return root

def get_all_topic_ids(topic_tree):
    """Retorna uma lista de todos os ids de tópicos na árvore."""
    ids = [topic_tree['id']]
    for child in topic_tree.get('children', []):
        ids.extend(get_all_topic_ids(child))
    return ids

def add_toc(paragraph):
    """Adiciona um campo de TOC (sumário) no docx."""
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)

def count_questions_in_subtree(topic_tree, questions_by_topic):
    """Conta o total de questões neste tópico e em todos os seus sub-tópicos."""
    total = len(questions_by_topic.get(topic_tree['id'], []))
    for child in topic_tree.get('children', []):
        total += count_questions_in_subtree(child, questions_by_topic)
    return total

def get_breadcrumb(topic_tree, numbering, parent_names=None):
    """Gera o breadcrumb do tópico atual, ex: 1. Obesidade > 1.1 Diagnóstico > 1.1.1 Avaliação Clínica"""
    if parent_names is None:
        parent_names = []
    breadcrumb_parts = []
    for i, (num, name) in enumerate(zip(numbering, parent_names + [topic_tree['nome']])):
        sub_numbering = '.'.join(str(n) for n in numbering[:i+1])
        breadcrumb_parts.append(f"{sub_numbering}. {name}")
    return ' > '.join(breadcrumb_parts)

def add_topic_sections_recursive(document, topic_tree, questions_by_topic, level=1, numbering=None, parent_names=None, questao_num=1, breadcrumb_raiz=None, permitir_repeticao=True, questoes_adicionadas=None, total_questoes_banco=1000):
    print(f"[LOG] Adicionando seção para tópico: {topic_tree['nome']} (ID: {topic_tree['id']})")
    
    # Usar o nível da árvore se disponível, senão usar o parâmetro level
    current_level = topic_tree.get('nivel', level)
    
    # Inicializar conjunto de questões adicionadas se não fornecido
    if questoes_adicionadas is None:
        questoes_adicionadas = set()
    
    # Verificar se o tópico tem questões antes de processá-lo
    total_questoes = count_questions_in_subtree(topic_tree, questions_by_topic)
    if total_questoes == 0:
        print(f"[LOG] Pulando tópico {topic_tree['nome']} - sem questões")
        # Processar apenas os filhos que têm questões
        for idx, child in enumerate(topic_tree.get('children', []), 1):
            print(f"[LOG] Verificando filho: {child['nome']} (ID: {child['id']})")
            questao_num = add_topic_sections_recursive(
                document,
                child,
                questions_by_topic,
                level=min(current_level+1, 9),
                numbering=numbering + [idx] if numbering else [1, idx],
                parent_names=parent_names + [topic_tree['nome']] if parent_names else [topic_tree['nome']],
                questao_num=questao_num,
                breadcrumb_raiz=breadcrumb_raiz,
                permitir_repeticao=permitir_repeticao,
                questoes_adicionadas=questoes_adicionadas,
                total_questoes_banco=total_questoes_banco
            )
        return questao_num
    
    if numbering is None:
        numbering = [1]
    else:
        numbering = numbering.copy()
    if parent_names is None:
        parent_names = []
    numbering_str = '.'.join(str(n) for n in numbering) + '.'
   
    # Calcular questões diretamente associadas ao tópico pai
    questoes_diretas = questions_by_topic.get(topic_tree['id'], [])
    total_questoes_filhos = total_questoes - len(questoes_diretas)
    
    heading_text = f"{numbering_str} {topic_tree['nome']} ({total_questoes} {'questões' if total_questoes != 1 else 'questão'})"

    # Variável para controlar se é o primeiro tópico de nível 1
    is_first_level1 = (current_level == 1 and numbering == [1])
    
    # Lógica de criação de seções baseada no número total de questões
    # Se <= 500 questões: apenas tópicos de nível 1 têm quebras de página
    # Se > 500 questões: tópicos de níveis 1, 2 e 3 têm quebras de página
    needs_new_section = False
    
    if total_questoes_banco <= 500:
        # Para bancos pequenos (<= 500): apenas nível 1 com quebra de página
        if current_level == 1 and not is_first_level1:
            needs_new_section = True
            print(f"[LOG] Banco pequeno ({total_questoes_banco} questões): quebra apenas nível 1")
    else:
        # Para bancos grandes (>= 500): níveis 1, 2 e 3 com quebra de página
        if current_level == 1 and not is_first_level1:
            # Criar nova seção para tópicos de nível 1 a partir do segundo
            needs_new_section = True
        elif current_level in [2, 3]:
            # Sempre criar nova seção para tópicos de nível 2 e 3
            needs_new_section = True
            print(f"[LOG] Banco grande ({total_questoes_banco} questões): quebra níveis 1-3")
    
    if needs_new_section:
        document.add_section(WD_SECTION.NEW_PAGE)
        print(f"[LOG] Nova seção criada para tópico nível {current_level}: {topic_tree['nome']}")
    
    # Adiciona breadcrumb no cabeçalho baseado no número de questões
    # Se <= 500: apenas nível 1 | Se >= 500: níveis 1, 2 e 3
    max_breadcrumb_level = 1 if total_questoes_banco <= 500 else 3
    if current_level <= max_breadcrumb_level:
        section = document.sections[-1]
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = True
        header = section.header
        for p in header.paragraphs:
            p.clear()
        
        # Gerar breadcrumb numerado para níveis 1, 2 e 3
        breadcrumb_parts = []
        
        # Construir lista com numerações e nomes dos ancestrais + tópico atual
        all_names = parent_names + [topic_tree['nome']]
        
        for i, name in enumerate(all_names):
            # Criar numeração parcial (ex: "1", "1.2", "1.2.3")
            partial_numbering = '.'.join(str(n) for n in numbering[:i+1])
            breadcrumb_parts.append(f"{partial_numbering}. {name}")
        
        breadcrumb_text = ' > '.join(breadcrumb_parts)
        print(f"[LOG] Breadcrumb criado para nível {current_level}: {breadcrumb_text}")
        
        p = header.paragraphs[0]
        p.clear()
        run = p.add_run(breadcrumb_text)
        run.bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.add_heading(heading_text, level=current_level)
    document.add_paragraph("")
    
    # Adiciona questões diretamente associadas ao tópico pai
    for q in questoes_diretas:
        # Verificar se a questão já foi adicionada (se não permitir repetição)
        if not permitir_repeticao and q['questao_id'] in questoes_adicionadas:
            print(f"[LOG] Pulando questão {q.get('codigo', '?')} - já adicionada anteriormente")
            continue
            
        print(f"[LOG] Adicionando questão {q.get('codigo', '?')} diretamente ao tópico {topic_tree['nome']}")
        
        # Adicionar questão ao conjunto de questões já adicionadas
        if not permitir_repeticao:
            questoes_adicionadas.add(q['questao_id'])
        
        # Determina o nível de dificuldade textual
        dificuldade_val = q.get('dificuldade', 0)
        try:
            dificuldade_val = int(dificuldade_val)
        except Exception:
            dificuldade_val = 0
        if dificuldade_val in [1, 2]:
            nivel_dificuldade = 'FÁCIL'
        elif dificuldade_val == 3:
            nivel_dificuldade = 'MÉDIO'
        elif dificuldade_val in [4, 5]:
            nivel_dificuldade = 'DIFÍCIL'
        else:
            nivel_dificuldade = ''
        # Monta o cabeçalho no padrão solicitado
        cabecalho = (
            f"{questao_num}. ({q['questao_id']}, {q['ano']}, {q.get('instituicao', '')}"
            f". Dificuldade: {nivel_dificuldade}). "
        )
        # Cria o parágrafo e adiciona o cabeçalho em negrito
        p = document.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY;
        run = p.add_run(clean_xml_illegal_chars(cabecalho))
        run.bold = True
        # Adiciona o enunciado (texto puro) na mesma linha
        enunciado_texto = extrair_texto_sem_imagens(q['enunciado'])
        p.add_run(clean_xml_illegal_chars(enunciado_texto))
        # Adiciona as imagens do enunciado (abaixo do texto)
        add_imagens_enunciado(document, q['enunciado'], q['codigo'], r"C:\Users\elman\OneDrive\Imagens\QuestoesResidencia")
        for alt in ['A', 'B', 'C', 'D', 'E']:
            alt_text = q.get(f'alternativa{alt}')
            if alt_text:
                safe_text = clean_xml_illegal_chars(f"{alt}) {alt_text}")
                document.add_paragraph(safe_text)
        document.add_paragraph("")
        p = document.add_paragraph()
        run = p.add_run("------  COMENTÁRIO  ------")
        run.bold = True
        run.font.color.rgb = RGBColor(0x1E, 0x90, 0xFF)
        p = document.add_paragraph()
        gabarito_texto_limpo = clean_xml_illegal_chars(q['gabarito_texto'])
        run = p.add_run(f"Gabarito: {q['gabarito']} - {gabarito_texto_limpo}")
        run.bold = True

        if q.get('gabaritoIA') == q.get('gabarito'):
            add_comentario_with_images(
                document,
                q['comentarioIA'],
                q['codigo'],
                r"C:\Users\elman\OneDrive\Imagens\QuestoesResidencia_comentarios",
                usar_src_absoluto=True
            )
        else:
            add_comentario_with_images(document, q['comentario'], q['codigo'], r"C:\Users\elman\OneDrive\Imagens\QuestoesResidencia_comentarios")
        document.add_paragraph("")  # Espaço
        questao_num += 1
    
    # Adiciona filhos recursivamente
    for idx, child in enumerate(topic_tree.get('children', []), 1):
        print(f"[LOG] Descendo para sub-tópico: {child['nome']} (ID: {child['id']})")
        questao_num = add_topic_sections_recursive(
            document,
            child,
            questions_by_topic,
            level=min(current_level+1, 9),
            numbering=numbering + [idx],
            parent_names=parent_names + [topic_tree['nome']],
            questao_num=questao_num,
            breadcrumb_raiz=breadcrumb_raiz,
            permitir_repeticao=permitir_repeticao,
            questoes_adicionadas=questoes_adicionadas,
            total_questoes_banco=total_questoes_banco
        )
    
    return questao_num

# Função para adicionar rodapé customizado em todas as seções
def add_footer_with_text_and_page_number(document):
    # Aplicar rodapé a todas as seções
    for section in document.sections:
        section.footer.is_linked_to_previous = False
        footer = section.footer
        # Limpa o rodapé existente
        for p in footer.paragraphs:
            p.clear()
        # Primeiro parágrafo: texto centralizado
        p_center = footer.add_paragraph()
        p_center.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p_center.add_run("Questões MED - 2025")
        # Segundo parágrafo: numeração de página à direita
        p_right = footer.add_paragraph()
        p_right.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run_right = p_right.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run_right._r.append(fldChar1)
        run_right._r.append(instrText)
        run_right._r.append(fldChar2)
        run_right._r.append(fldChar3)
        p_right.add_run(" de ")
        run_total = p_right.add_run()
        fldChar1t = OxmlElement('w:fldChar')
        fldChar1t.set(qn('w:fldCharType'), 'begin')
        instrTextt = OxmlElement('w:instrText')
        instrTextt.text = 'NUMPAGES'
        fldChar2t = OxmlElement('w:fldChar')
        fldChar2t.set(qn('w:fldCharType'), 'separate')
        fldChar3t = OxmlElement('w:fldChar')
        fldChar3t.set(qn('w:fldCharType'), 'end')
        run_total._r.append(fldChar1t)
        run_total._r.append(instrTextt)
        run_total._r.append(fldChar2t)
        run_total._r.append(fldChar3t)

def render_mermaid_to_image(mermaid_code, temp_dir):
    """
    Renderiza código Mermaid para PNG usando API externa.
    Retorna o caminho da imagem gerada ou None se falhar.
    """
    import hashlib
    import base64
    
    try:
        # Criar hash do código Mermaid para usar como nome do arquivo
        mermaid_hash = hashlib.md5(mermaid_code.encode('utf-8')).hexdigest()[:12]
        output_path = os.path.join(temp_dir, f"mermaid_{mermaid_hash}.png")
        
        # Se a imagem já existe, retornar o caminho
        if os.path.exists(output_path):
            print(f"[LOG] Imagem Mermaid já existe: {output_path}")
            return output_path
        
        # Codificar o código Mermaid em base64 para a API
        mermaid_encoded = base64.urlsafe_b64encode(mermaid_code.encode('utf-8')).decode('ascii')
        
        # URL da API do Mermaid (serviço público de renderização)
        api_url = f"https://mermaid.ink/img/{mermaid_encoded}"
        
        print(f"[LOG] Renderizando Mermaid via API: {api_url[:80]}...")
        
        # Fazer requisição para a API
        response = requests.get(api_url, timeout=30)
        
        if response.status_code == 200:
            # Salvar a imagem
            os.makedirs(temp_dir, exist_ok=True)
            with open(output_path, 'wb') as f:
                f.write(response.content)
            print(f"[LOG] Imagem Mermaid gerada: {output_path}")
            return output_path
        else:
            print(f"[AVISO] Falha ao renderizar Mermaid: status {response.status_code}")
            return None
            
    except Exception as e:
        print(f"[AVISO] Erro ao renderizar Mermaid: {str(e)}")
        return None

def add_comentario_with_images(document, comentario_md, codigo_questao, imagens_dir, usar_src_absoluto=False):
    # Reduz múltiplas linhas em branco para apenas uma (\n\n), mantendo parágrafos separados
    comentario_md = re.sub(r'\n{3,}', '\n\n', comentario_md)
    
    # Processar blocos Mermaid antes de converter para HTML
    # Criar diretório temporário para imagens Mermaid
    temp_mermaid_dir = os.path.join(os.path.dirname(__file__), 'temp_mermaid')
    os.makedirs(temp_mermaid_dir, exist_ok=True)
    
    # Regex para encontrar blocos de código Mermaid
    mermaid_pattern = r'```mermaid\s*\n(.*?)```'
    mermaid_blocks = re.finditer(mermaid_pattern, comentario_md, re.DOTALL)
    
    # Armazenar imagens Mermaid processadas
    mermaid_images_map = {}
    
    # Processar cada bloco Mermaid encontrado (processar em ordem reversa para manter posições corretas)
    mermaid_blocks = list(mermaid_blocks)
    for idx, match in enumerate(reversed(mermaid_blocks)):
        mermaid_code = match.group(1).strip()
        print(f"[LOG] Processando bloco Mermaid {len(mermaid_blocks) - idx}")
        
        # Renderizar Mermaid para imagem
        img_path = render_mermaid_to_image(mermaid_code, temp_mermaid_dir)
        
        if img_path and os.path.exists(img_path):
            # Criar uma tag HTML img inline que será processada pelo parser HTML
            # Usar um placeholder temporário que será substituído
            placeholder = f'<img src="MERMAID_TEMP_{len(mermaid_blocks) - idx - 1}" />'
            comentario_md = comentario_md[:match.start()] + placeholder + comentario_md[match.end():]
            
            # Armazenar o caminho da imagem
            mermaid_images_map[len(mermaid_blocks) - idx - 1] = img_path
    
    html = markdown(comentario_md, extras=['tables'])
    
    # Substituir placeholders MERMAID_TEMP_ pelos caminhos reais
    for idx, img_path in mermaid_images_map.items():
        html = html.replace(f'src="MERMAID_TEMP_{idx}"', f'src="{img_path}"')
    soup = BeautifulSoup(html, "html.parser")
    img_count = [1]

    def add_horizontal_rule():
        """Adiciona uma linha divisória horizontal no documento"""
        # Criar parágrafo para a linha divisória
        p = document.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Adicionar linha de caracteres para simular linha divisória
        run = p.add_run("─" * 50)  # 50 caracteres de linha
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)  # Cinza
        run.font.size = Pt(10)
        
        # Adicionar espaçamento antes e depois
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

    def add_table_from_html(table_element):
        """Converte uma tabela HTML para uma tabela DOCX"""
        # Encontrar todas as linhas (tr)
        rows = table_element.find_all('tr')
        if not rows:
            return
        
        # Determinar número de colunas
        max_cols = 0
        for row in rows:
            cells = row.find_all(['th', 'td'])
            max_cols = max(max_cols, len(cells))
        
        if max_cols == 0:
            return
        
        # Criar tabela no DOCX
        table = document.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Table Grid'  # Estilo com bordas
        
        # Preencher tabela
        for row_idx, row in enumerate(rows):
            cells = row.find_all(['th', 'td'])
            docx_row = table.rows[row_idx]
            
            for col_idx, cell in enumerate(cells):
                if col_idx < max_cols:
                    # Obter texto da célula
                    cell_text = cell.get_text().strip()
                    cell_text = clean_xml_illegal_chars(cell_text)
                    
                    # Adicionar texto à célula do DOCX
                    docx_cell = docx_row.cells[col_idx]
                    docx_paragraph = docx_cell.paragraphs[0]
                    
                    # Verificar se é cabeçalho (th) e aplicar formatação
                    if cell.name == 'th':
                        # Cabeçalho: negrito e fundo cinza claro
                        run = docx_paragraph.add_run(cell_text)
                        run.bold = True
                        # Tentar aplicar fundo cinza (pode não funcionar em todas as versões)
                        try:
                            docx_cell._tc.get_or_add_tcPr().append(
                                OxmlElement('w:shd')
                            ).set(qn('w:fill'), 'D9D9D9')
                        except:
                            pass  # Se não conseguir aplicar cor de fundo, continua
                    else:
                        # Célula normal
                        docx_paragraph.add_run(cell_text)
                    
                    # Centralizar texto nas células
                    docx_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Adicionar espaçamento após a tabela
        document.add_paragraph("")

    def add_heading_from_html(heading_element):
        """Converte um título HTML para um parágrafo formatado (não usa estilo de título para evitar aparecer no sumário)"""
        # Determinar nível do título (h1 = 1, h2 = 2, etc.)
        level = int(heading_element.name[1])  # Remove 'h' e converte para int
        
        # Obter texto do título
        heading_text = heading_element.get_text().strip()
        heading_text = clean_xml_illegal_chars(heading_text)
        
        if not heading_text:
            return
        
        # Mapear níveis de título para tamanhos de fonte
        # h1 -> 14pt, h2 -> 13pt, h3 -> 12pt, h4+ -> 11pt
        font_size_mapping = {
            1: Pt(14),  # h1
            2: Pt(13),  # h2
            3: Pt(12),  # h3
        }
        
        # Usar tamanho 11pt para níveis 4 e superiores
        font_size = font_size_mapping.get(level, Pt(11))
        
        # Adicionar como parágrafo normal com formatação especial
        p = document.add_paragraph()
        run = p.add_run(heading_text)
        run.bold = True
        run.font.size = font_size

    def add_formatted_paragraph(text, level=0, is_bullet=False, bullet_char="•"):
        """Adiciona um parágrafo formatado com indentação e formatação"""
        if not text or not text.strip():
            return
            
        # Limpar texto
        text = clean_xml_illegal_chars(text.strip())
        
        # Criar parágrafo
        p = document.add_paragraph()
        
        # Aplicar indentação baseada no nível
        if level > 0:
            p.paragraph_format.left_indent = Inches(0.25 * level)
        
        # Adicionar bullet se necessário
        if is_bullet:
            run = p.add_run(f"{bullet_char} ")
            run.bold = True
        
        # Processar texto com formatação (negrito, etc.)
        add_formatted_text(p, text)

    def add_formatted_text(paragraph, text):
        """Adiciona texto com formatação (negrito, etc.)"""
        # Processar texto em negrito (**texto**)
        parts = re.split(r'(\*\*.*?\*\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Texto em negrito
                bold_text = part[2:-2]  # Remove **
                run = paragraph.add_run(bold_text)
                run.bold = True
            else:
                # Texto normal - manter espaços e quebras de linha
                if part:  # Não usar .strip() para preservar espaços
                    paragraph.add_run(part)

    def process_list_item(li_element, level=0):
        """Processa um item de lista com indentação apropriada"""
        # Coletar texto do item (apenas texto direto, não sublistas)
        item_text = []
        
        for child in li_element.children:
            if isinstance(child, str):
                text = child.replace('\xa0', ' ').strip()
                if text:
                    item_text.append(text)
            elif child.name in ['strong', 'b']:
                # Texto em negrito
                bold_text = child.get_text().strip()
                if bold_text:
                    item_text.append(f"**{bold_text}**")
            elif child.name in ['em', 'i']:
                # Texto em itálico
                italic_text = child.get_text().strip()
                if italic_text:
                    item_text.append(f"*{italic_text}*")
            elif child.name not in ['ul', 'ol']:
                # Outros elementos (exceto listas)
                text = child.get_text().strip()
                if text:
                    item_text.append(text)
        
        # Adicionar o item da lista (se houver texto)
        if item_text:
            full_text = ' '.join(item_text)
            bullet_char = "•" if level == 0 else "▪" if level == 1 else "▫"
            add_formatted_paragraph(full_text, level, is_bullet=True, bullet_char=bullet_char)
        
        # Processar sublistas APÓS o texto do item
        for child in li_element.children:
            if child.name in ['ul', 'ol']:
                process_list(child, level + 1)

    def process_list(list_element, level=0):
        """Processa uma lista (ul ou ol)"""
        # Processar apenas itens diretos (não recursivos)
        for li in list_element.find_all('li', recursive=False):
            process_list_item(li, level)

    def obter_caminho_imagem(src, indice_imagem):
        """
        Retorna o caminho da imagem a ser inserida.
        Se usar_src_absoluto=True e o src apontar para um arquivo existente, usa-o diretamente.
        Caso contrário, monta o caminho usando imagens_dir e o padrão legado.
        """
        if usar_src_absoluto and src:
            # Normalizar barras para evitar problemas no Windows
            src_normalizado = os.path.normpath(src)
            if os.path.isabs(src_normalizado) and os.path.exists(src_normalizado):
                return src_normalizado
            # Tentar relativo ao diretório atual do script
            relativo_script = os.path.join(os.path.dirname(__file__), src_normalizado)
            if os.path.exists(relativo_script):
                return relativo_script
        # Fallback para comportamento original
        ext = os.path.splitext(src)[1].split("?")[0]
        if not ext:
            ext = ".jpeg"
        if indice_imagem == 1:
            img_filename = f"{codigo_questao}{ext}"
        else:
            img_filename = f"{codigo_questao}_{indice_imagem}{ext}"
        return os.path.join(imagens_dir, img_filename)

    def process_element(elem):
        if isinstance(elem, Comment):
            return
        elif isinstance(elem, str):
            text = elem.replace('\xa0', ' ').strip()
            if text:
                add_formatted_paragraph(text)
        elif elem.name == "img":
            src = elem.get("src", "")
            img_path = obter_caminho_imagem(src, img_count[0])
            max_width = get_max_image_width(document)
            if not verificar_e_adicionar_imagem(document, img_path, max_width):
                document.add_paragraph(f"[Imagem não encontrada ou inválida: {os.path.basename(img_path)}]")
            img_count[0] += 1
        elif elem.name in ["br"]:
            # Quebra de linha
            document.add_paragraph("")
        elif elem.name == "hr":
            # Linha divisória horizontal (---)
            add_horizontal_rule()
        elif elem.name == "table":
            # Tabela HTML - converter para tabela DOCX
            add_table_from_html(elem)
        elif elem.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            # Título HTML - converter para título DOCX
            add_heading_from_html(elem)
        elif elem.name in ["div", "p"]:
            # CORREÇÃO: Coletar todo o texto do parágrafo antes de processar
            # para evitar quebras de linha desnecessárias
            paragraph_text = []
            for child in elem.children:
                if hasattr(child, 'name'):
                    if child.name == "img":
                        # Processar imagem diretamente
                        src = child.get("src", "")
                        img_path = obter_caminho_imagem(src, img_count[0])
                        max_width = get_max_image_width(document)
                        if not verificar_e_adicionar_imagem(document, img_path, max_width):
                            document.add_paragraph(f"[Imagem não encontrada ou inválida: {os.path.basename(img_path)}]")
                        img_count[0] += 1
                    elif child.name in ["strong", "b"]:
                        # Texto em negrito
                        bold_text = child.get_text().strip()
                        if bold_text:
                            paragraph_text.append(f"**{bold_text}**")
                    elif child.name in ["em", "i"]:
                        # Texto em itálico
                        italic_text = child.get_text().strip()
                        if italic_text:
                            paragraph_text.append(f"*{italic_text}*")
                    elif child.name not in ["ul", "ol"]:
                        # Outros elementos (exceto listas)
                        text = child.get_text().strip()
                        if text:
                            paragraph_text.append(text)
                elif isinstance(child, str):
                    text = child.replace('\xa0', ' ').strip()
                    if text:
                        paragraph_text.append(text)
            
            # Adicionar como um único parágrafo se houver conteúdo
            if paragraph_text:
                full_text = ' '.join(paragraph_text)
                add_formatted_paragraph(full_text)
        elif elem.name in ["ul", "ol"]:
            # Processar lista
            process_list(elem, 0)
        elif elem.name == "strong" or elem.name == "b":
            # Texto em negrito (fallback para elementos isolados)
            bold_text = elem.get_text().strip()
            if bold_text:
                add_formatted_paragraph(f"**{bold_text}**")
        elif elem.name == "em" or elem.name == "i":
            # Texto em itálico (fallback para elementos isolados)
            italic_text = elem.get_text().strip()
            if italic_text:
                add_formatted_paragraph(f"*{italic_text}*")
        else:
            # Outros elementos - processar filhos
            for child in elem.children:
                process_element(child)

    # Processar todos os elementos
    for elem in soup.contents:
        process_element(elem)

def get_max_image_width(document):
    """
    Calcula a largura máxima disponível para imagens na página.
    Esta largura é usada como LIMITE MÁXIMO, não como largura forçada.
    """
    section = document.sections[-1]
    page_width = section.page_width
    left_margin = section.left_margin
    right_margin = section.right_margin
    return page_width - left_margin - right_margin

# Função para extrair apenas o texto do enunciado, sem imagens
def extrair_texto_sem_imagens(enunciado_html):
    soup = BeautifulSoup(enunciado_html, "html.parser")
    for img in soup.find_all('img'):
        img.decompose()
    return soup.get_text(separator=" ").replace('\xa0', ' ').strip()

# Função para adicionar apenas as imagens do enunciado
def add_imagens_enunciado(document, enunciado_html, codigo_questao, imagens_dir):
    soup = BeautifulSoup(enunciado_html, "html.parser")
    img_count = 1
    for img in soup.find_all('img'):
        if img_count == 1:
            img_filename = f"{codigo_questao}.jpeg"
        else:
            img_filename = f"{codigo_questao}_{img_count}.jpeg"
        img_path = os.path.join(imagens_dir, img_filename)
        max_width = get_max_image_width(document)
        if not verificar_e_adicionar_imagem(document, img_path, max_width):
            document.add_paragraph(f"[Imagem não encontrada ou inválida: {img_filename}]")
        img_count += 1

def clean_xml_illegal_chars(text):
    # Remove caracteres de controle e inválidos para XML (exceto \t, \n, \r)
    # Inclui \ufffe, \uffff, e outros fora do intervalo permitido
    illegal_unichrs = [
        (0x00, 0x08), (0x0B, 0x0C), (0x0E, 0x1F),
        (0x7F, 0x84), (0x86, 0x9F),
        (0xFDD0, 0xFDDF), (0xFFFE, 0xFFFF)
    ]
    re_illegal = u'|'.join('%s-%s' % (chr(low), chr(high)) for (low, high) in illegal_unichrs)
    re_illegal = '[%s]' % re_illegal
    text = re.sub(re_illegal, '', text)
    # Remove qualquer outro caractere de controle ASCII, exceto \t, \n, \r
    text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
    return text

def limpar_nome_para_titulo(nome):
    """
    Remove caracteres inválidos e normaliza espaços para uso em metadados.
    """
    if not nome:
        return "Banco de Questões"
    nome_limpo = str(nome).strip()
    nome_limpo = re.sub(r'[\\/:*?"<>|]+', ' ', nome_limpo)
    nome_limpo = re.sub(r'\s+', ' ', nome_limpo)
    return nome_limpo or "Banco de Questões"

def configurar_metadados_documento(document, total_questoes, nome_arquivo_limpo=None):
    """
    Configura os metadados do documento DOCX.
    
    Args:
        document: Objeto Document do python-docx
        total_questoes: Número total de questões no banco
        nome_arquivo_limpo: Nome amigável para composição do título
    """
    print("[LOG] Configurando metadados do documento...")
    
    # === PROPRIEDADES PRINCIPAIS (CORE PROPERTIES) ===
    
    # 👤 Autor do documento
    document.core_properties.author = "Questões MED"
    
    # 📝 Título do documento
    titulo_base = limpar_nome_para_titulo(nome_arquivo_limpo)
    document.core_properties.title = f"E-book de Questões Comentadas de {titulo_base}"
    
    # 📚 Assunto/Tema
    document.core_properties.subject = "Banco de Questões de Medicina"
    
    # 🏷️ Palavras-chave (separadas por vírgula)
    document.core_properties.keywords = "medicina, residência médica, banco de questões"
    
    # 👔 Gerente/Responsável
    document.core_properties.manager = "Questões MED"
    
    # 📂 Categoria
    document.core_properties.category = "Educação Médica"
    
    # 💭 Comentários/Descrição
    data_geracao = datetime.now().strftime('%d/%m/%Y às %H:%M')
    document.core_properties.comments = (
        f"Banco de questões de provas de residência médica"
        f"Contém questões comentadas e organizadas em uma ampla hierarquia de tópicos."
    )
    
    # 📊 Último editor
    document.core_properties.last_modified_by = "Sistema Gerador de Banco de Questões MED"
    
    # 🔄 Número da revisão
    document.core_properties.revision = 1
    
    # ⏰ Data de criação
    document.core_properties.created = datetime.now()
    
    # 📅 Data de modificação
    document.core_properties.modified = datetime.now()
    
    print(f"[LOG] Metadados configurados:")
    print(f"  - Autor: {document.core_properties.author}")
    print(f"  - Título: {document.core_properties.title}")
    print(f"  - Categoria: {document.core_properties.category}")
    print(f"  - Palavras-chave: {document.core_properties.keywords}")
    print(f"  - Data criação: {document.core_properties.created.strftime('%d/%m/%Y %H:%M')}")

def gerar_banco_estratificacao_deterministica(conn, total_questoes=1000, permitir_repeticao=True):
    """
    Gera um banco de questões usando consulta SQL específica com N questões
    e organizando hierarquicamente com profundidade máxima de nível 4.
    """
    print(f"[LOG] Gerando banco de questões com consulta SQL específica - {total_questoes} questões...")
    
    # Informar comportamento de seções baseado no número de questões
    if total_questoes <= 500:
        print(f"[LOG] Banco COMPACTO ({total_questoes} questões): quebras de página apenas para tópicos de NÍVEL 1")
    else:
        print(f"[LOG] Banco EXPANDIDO ({total_questoes} questões): quebras de página para tópicos de NÍVEIS 1, 2 e 3")
    
    # Executar a consulta SQL fornecida para obter as questões selecionadas
    cursor = conn.cursor(dictionary=True)
    
    # Calcular cotas por área baseado no total N (ordem: 1.Cirurgia, 2.Clínica Médica, 3.Pediatria, 4.Ginecologia, 5.Obstetrícia, 6.Med.Preventiva)
    cotas = {
        'Cirurgia': round(total_questoes * 0.2),
        'Clínica Médica': round(total_questoes * 0.2),
        'Pediatria': round(total_questoes * 0.2),
        'Ginecologia': round(total_questoes * 0.1),
        'Obstetrícia': round(total_questoes * 0.1),
        'Medicina Preventiva': round(total_questoes * 0.2)
    }
    
    print(f"[LOG] Cotas calculadas para {total_questoes} questões: {cotas}")
    print(f"[LOG] Usando consulta com tópicos raiz específicos e ordenação SHA2 determinística")
    
    query_questoes = f"""
    WITH cotas AS (
        SELECT 33   AS topico_id_raiz, 'Cirurgia'            AS area, ROUND({total_questoes} * 0.20) AS qtd
        UNION ALL
        SELECT 100  AS topico_id_raiz, 'Clínica Médica'      AS area, ROUND({total_questoes} * 0.20)
        UNION ALL
        SELECT 48   AS topico_id_raiz, 'Pediatria'           AS area, ROUND({total_questoes} * 0.20)
        UNION ALL
        SELECT 183  AS topico_id_raiz, 'Ginecologia'         AS area, ROUND({total_questoes} * 0.10)
        UNION ALL
        SELECT 218  AS topico_id_raiz, 'Obstetrícia'         AS area, ROUND({total_questoes} * 0.10)
        UNION ALL
        SELECT 29   AS topico_id_raiz, 'Medicina Preventiva' AS area, ROUND({total_questoes} * 0.20)
    ),
    ordenadas AS (
        SELECT 
            q.*,
            ROW_NUMBER() OVER (
                PARTITION BY c.area
                ORDER BY SHA2(CONCAT(q.questao_id, 'SEMENTE_FIXA'), 256)
            ) AS ordem,
            c.qtd
        FROM questaoresidencia q
        JOIN cotas c ON q.area = c.area
        WHERE CHAR_LENGTH(q.comentario) >= 500 AND q.ano >= 2018
    )
    SELECT 
        o.*
    FROM ordenadas o
    WHERE o.ordem <= o.qtd
    ORDER BY o.area, o.ordem
    """
    
    print("[LOG] Executando consulta SQL para selecionar questões...")
    cursor.execute(query_questoes)
    questoes_selecionadas = cursor.fetchall()
    
    print(f"[LOG] Total de questões selecionadas: {len(questoes_selecionadas)}")
    
    # Mostrar distribuição por área das questões selecionadas
    distribuicao_selecionadas = {}
    for q in questoes_selecionadas:
        area = q['area']
        distribuicao_selecionadas[area] = distribuicao_selecionadas.get(area, 0) + 1
    
    print("[LOG] Distribuição por área das questões selecionadas:")
    for area, count in distribuicao_selecionadas.items():
        print(f"  - {area}: {count} questões")
    
    # Mapear áreas para tópicos raiz conforme definido na consulta (ordem: 1.Cirurgia, 2.Clínica Médica, 3.Pediatria, 4.Ginecologia, 5.Obstetrícia, 6.Med.Preventiva)
    area_para_topico_raiz = {
        'Cirurgia': 33,
        'Clínica Médica': 100,
        'Pediatria': 48,
        'Ginecologia': 183,
        'Obstetrícia': 218,
        'Medicina Preventiva': 29
    }
    
    print(f"[LOG] Mapeamento área -> tópico raiz: {area_para_topico_raiz}")
    
    # Associar cada questão ao seu tópico raiz baseado na área
    questoes_sem_topico = 0
    for q in questoes_selecionadas:
        area = q['area']
        topico_raiz = area_para_topico_raiz.get(area)
        if topico_raiz:
            q['id_topico'] = topico_raiz
        else:
            print(f"[ERRO] Área '{area}' não mapeada para tópico raiz")
            q['id_topico'] = None
            questoes_sem_topico += 1
    
    if questoes_sem_topico == 0:
        print(f"[LOG] Todas as questões associadas aos tópicos raiz por área")
    else:
        print(f"[ERRO] {questoes_sem_topico} questões não puderam ser associadas a tópicos")
    
    # Obter questões com classificações mais específicas para melhor organização
    questao_ids = [q['questao_id'] for q in questoes_selecionadas]
    format_strings = ','.join(['%s'] * len(questao_ids))
    
    query_topicos_especificos = f"""
    SELECT DISTINCT cq.id_topico, cq.id_questao
    FROM classificacao_questao cq
    WHERE cq.id_questao IN ({format_strings})
    ORDER BY cq.id_questao, cq.id_topico
    """
    
    cursor.execute(query_topicos_especificos, tuple(questao_ids))
    classificacoes_especificas = cursor.fetchall()
    
    print(f"[LOG] Classificações específicas encontradas: {len(classificacoes_especificas)}")
    
    # Criar mapeamento de questão -> tópicos específicos para melhor organização
    questao_topicos_especificos = {}
    for classificacao in classificacoes_especificas:
        questao_id = classificacao['id_questao']
        topico_id = classificacao['id_topico']
        if questao_id not in questao_topicos_especificos:
            questao_topicos_especificos[questao_id] = []
        questao_topicos_especificos[questao_id].append(topico_id)
    
    # Usar tópico mais específico se disponível, senão manter tópico raiz
    for q in questoes_selecionadas:
        topicos_especificos = questao_topicos_especificos.get(q['questao_id'], [])
        if topicos_especificos:
            # Usar o primeiro tópico específico encontrado para melhor organização
            q['id_topico'] = topicos_especificos[0]
        # Se não houver tópico específico, mantém o tópico raiz já definido
    
    # Como usamos INNER JOIN, todas as questões têm tópico associado
    questoes_com_topico = questoes_selecionadas
    print(f"[LOG] Questões com tópico associado: {len(questoes_com_topico)}")
    
    # Verificar se obtivemos exatamente o número esperado
    if len(questoes_com_topico) < total_questoes:
        diferenca = total_questoes - len(questoes_com_topico)
        print(f"[AVISO] Obtidas apenas {len(questoes_com_topico)} questões de {total_questoes} solicitadas.")
        print(f"[AVISO] Diferença: {diferenca} questões. Isso pode indicar que não há questões suficientes")
        print(f"[AVISO] no banco que atendam aos critérios (comentário ≥500 chars, ano ≥2018, etc.)")
    
    # Mostrar distribuição final por área
    distribuicao_final = {}
    for q in questoes_com_topico:
        area = q['area']
        distribuicao_final[area] = distribuicao_final.get(area, 0) + 1
    
    print("[LOG] Distribuição final por área:")
    for area, count in distribuicao_final.items():
        cota_esperada = cotas.get(area, 0)
        status = "✅" if count == cota_esperada else f"❌ (esperado: {cota_esperada})"
        print(f"  - {area}: {count} questões {status}")
    
    if len(questoes_com_topico) == total_questoes:
        print(f"✅ [SUCESSO] Exatamente {total_questoes} questões obtidas!")
    else:
        print(f"⚠️ [AVISO] Obtidas {len(questoes_com_topico)} questões de {total_questoes} solicitadas")
    
    # Obter todos os tópicos únicos das questões
    topicos_utilizados = list(set([q['id_topico'] for q in questoes_com_topico]))
    print(f"[LOG] Tópicos únicos utilizados: {len(topicos_utilizados)}")
    
    # Organizar questões por tópico
    questions_by_topic = {}
    for q in questoes_com_topico:
        tid = q['id_topico']
        if tid not in questions_by_topic:
            questions_by_topic[tid] = []
        questions_by_topic[tid].append(q)
    
    print(f"[LOG] Questões organizadas por {len(questions_by_topic)} tópicos")
    
    # Construir hierarquia completa dos tópicos utilizados
    print("[LOG] Construindo hierarquia completa dos tópicos...")
    
    # Obter hierarquia completa dos tópicos (incluindo ancestrais)
    topicos_completos = set(topicos_utilizados)
    
    # Para cada tópico utilizado, buscar todos os ancestrais
    for topico_id in topicos_utilizados:
        cursor.execute("""
            WITH RECURSIVE topico_ancestrais AS (
                SELECT id, id_pai, nome, 1 as nivel
                FROM topico 
                WHERE id = %s
                
                UNION ALL
                
                SELECT t.id, t.id_pai, t.nome, ta.nivel + 1
                FROM topico t
                INNER JOIN topico_ancestrais ta ON ta.id_pai = t.id
                WHERE ta.nivel < 10
            )
            SELECT id FROM topico_ancestrais
        """, (topico_id,))
        
        ancestrais = cursor.fetchall()
        for ancestral in ancestrais:
            topicos_completos.add(ancestral['id'])
    
    print(f"[LOG] Tópicos completos (incluindo ancestrais): {len(topicos_completos)}")
    
    # Buscar informações completas dos tópicos
    topicos_completos_list = list(topicos_completos)
    format_strings = ','.join(['%s'] * len(topicos_completos_list))
    
    cursor.execute(f"""
        SELECT id, nome, id_pai
        FROM topico 
        WHERE id IN ({format_strings})
        ORDER BY id
    """, tuple(topicos_completos_list))
    
    topicos_info = {t['id']: t for t in cursor.fetchall()}
    
    # Construir árvores hierárquicas
    def build_topic_tree(topico_id, nivel_atual=1, max_nivel=4):
        if topico_id not in topicos_info:
            return None
            
        topico = topicos_info[topico_id]
        tree_node = {
            'id': topico_id,
            'nome': topico['nome'],
            'nivel': nivel_atual,
            'children': []
        }
        
        # Se já atingiu o nível máximo, não adicionar mais filhos na árvore
        # mas as questões dos descendentes serão reagrupadas neste nível
        if nivel_atual >= max_nivel:
            return tree_node
        
        # Encontrar filhos diretos
        filhos = [t_id for t_id, t_info in topicos_info.items() 
                 if t_info['id_pai'] == topico_id and t_id in topicos_completos]
        
        for filho_id in sorted(filhos):
            child_tree = build_topic_tree(filho_id, nivel_atual + 1, max_nivel)
            if child_tree:
                tree_node['children'].append(child_tree)
        
        return tree_node
    
    # Encontrar tópicos raiz (sem pai ou pai não está no conjunto)
    topicos_raiz = []
    for topico_id in topicos_completos:
        topico = topicos_info[topico_id]
        if topico['id_pai'] is None or topico['id_pai'] not in topicos_completos:
            topicos_raiz.append(topico_id)
    
    print(f"[LOG] Tópicos raiz encontrados: {len(topicos_raiz)}")
    
    # Construir árvores para cada tópico raiz
    topic_trees = []
    for raiz_id in sorted(topicos_raiz):
        tree = build_topic_tree(raiz_id)
        if tree:
            topic_trees.append(tree)
    
    print(f"[LOG] Árvores construídas: {len(topic_trees)}")
    
    # Definir ordem específica das áreas médicas conforme solicitado
    ordem_areas = [
        'Cirurgia',
        'Clínica Médica',
        'Pediatria', 
        'Ginecologia',
        'Obstetrícia',
        'Medicina Preventiva'
    ]
    
    # Função para determinar a área de um tópico baseado nas questões
    def get_area_from_topic(tree, questions_by_topic):
        # Buscar questões do tópico e seus filhos para determinar a área
        def collect_questions_from_tree(node):
            all_questions = []
            if node['id'] in questions_by_topic:
                all_questions.extend(questions_by_topic[node['id']])
            for child in node.get('children', []):
                all_questions.extend(collect_questions_from_tree(child))
            return all_questions
        
        questoes = collect_questions_from_tree(tree)
        if questoes:
            # Usar a área da primeira questão como representativa
            return questoes[0].get('area', 'Outros')
        return 'Outros'
    
    # Organizar árvores por área
    arvores_por_area = {}
    for tree in topic_trees:
        area = get_area_from_topic(tree, questions_by_topic)
        if area not in arvores_por_area:
            arvores_por_area[area] = []
        arvores_por_area[area].append(tree)
    
    print(f"[LOG] Árvores organizadas por área: {list(arvores_por_area.keys())}")
    
    # Ordenar árvores conforme a sequência desejada: 1.Cirurgia, 2.Clínica Médica, 3.Pediatria, 4.Ginecologia, 5.Obstetrícia, 6.Med.Preventiva
    topic_trees_ordenadas = []
    for i, area in enumerate(ordem_areas, 1):
        if area in arvores_por_area:
            # Ordenar árvores da mesma área por nome do tópico
            arvores_area = sorted(arvores_por_area[area], key=lambda x: x['nome'])
            topic_trees_ordenadas.extend(arvores_area)
            print(f"[LOG] {i}. Adicionada área '{area}' com {len(arvores_area)} árvore(s)")
    
    # Adicionar áreas não mapeadas no final
    for area, arvores in arvores_por_area.items():
        if area not in ordem_areas:
            arvores_area = sorted(arvores, key=lambda x: x['nome'])
            topic_trees_ordenadas.extend(arvores_area)
            print(f"[LOG] Adicionada área adicional '{area}' com {len(arvores_area)} árvore(s)")
    
    topic_trees = topic_trees_ordenadas
    print(f"[LOG] Árvores reordenadas conforme sequência solicitada: {len(topic_trees)} árvores")
    
    # Reorganizar questões para tópicos de nível 4 (agrupar descendentes)
    def get_all_descendants(topico_id):
        """Retorna todos os descendentes de um tópico (incluindo ele próprio)"""
        descendants = {topico_id}
        
        # Buscar filhos diretos
        filhos = [t_id for t_id, t_info in topicos_info.items() 
                 if t_info['id_pai'] == topico_id]
        
        for filho_id in filhos:
            descendants.update(get_all_descendants(filho_id))
        
        return descendants
    
    def reorganize_questions_for_level4(tree_node, questions_by_topic, reorganized_questions):
        """Reorganiza questões para que tópicos de nível 4 incluam questões de todos os descendentes"""
        
        if tree_node['nivel'] == 4:
            # Este é um tópico de nível 4, coletar questões de todos os descendentes
            all_descendants = get_all_descendants(tree_node['id'])
            todas_questoes = []
            questoes_ids_unicos = set()  # Para evitar duplicatas
            
            for desc_id in all_descendants:
                if desc_id in questions_by_topic:
                    for questao in questions_by_topic[desc_id]:
                        # Verificar se a questão já foi adicionada (evitar duplicatas)
                        if questao['questao_id'] not in questoes_ids_unicos:
                            todas_questoes.append(questao)
                            questoes_ids_unicos.add(questao['questao_id'])
            
            if todas_questoes:
                reorganized_questions[tree_node['id']] = todas_questoes
                print(f"[LOG] Tópico nível 4 '{tree_node['nome']}': {len(todas_questoes)} questões reagrupadas (duplicatas removidas)")
            
        elif tree_node['nivel'] < 4:
            # Para níveis menores que 4, manter questões diretas e processar filhos
            if tree_node['id'] in questions_by_topic:
                reorganized_questions[tree_node['id']] = questions_by_topic[tree_node['id']]
            
            # Processar filhos recursivamente
            for child in tree_node['children']:
                reorganize_questions_for_level4(child, questions_by_topic, reorganized_questions)
    
    # Aplicar reorganização
    reorganized_questions = {}
    for tree in topic_trees:
        reorganize_questions_for_level4(tree, questions_by_topic, reorganized_questions)
    
    print(f"[LOG] Questões reorganizadas para {len(reorganized_questions)} tópicos")
    
    # Criar documento
    document = Document()
    
    # Configurar metadados do documento
    nome_titulo_documento = f"{total_questoes} Questões"
    configurar_metadados_documento(document, total_questoes, nome_titulo_documento)
    
    # Configurar estilo padrão
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(3)
    paragraph_format.space_before = Pt(0)
    paragraph_format.line_spacing = 1
    
    # === SEÇÃO 1: CAPA ===
    print("[LOG] Criando seção da capa...")
    
    # Configurar cabeçalho da capa com logotipo
    section_capa = document.sections[0]
    section_capa.header.is_linked_to_previous = False
    header_capa = section_capa.header
    for p in header_capa.paragraphs:
        p.clear()
    
    # Adicionar logotipo no cabeçalho (se disponível)
    img_path = os.path.join(os.path.dirname(__file__), 'img', 'logotipo.png')
    p_header = header_capa.paragraphs[0]
    p_header.clear()
    p_header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    logotipo_adicionado = False
    if os.path.exists(img_path):
        print(f"[LOG] Verificando logotipo: {img_path}")
        run_header = p_header.add_run()
        try:
            # Verificar se é uma imagem válida tentando abrir com PIL
            Image.open(img_path).verify()  # Verificar se é uma imagem válida
            
            run_header.add_picture(img_path, width=Inches(3))
            print(f"[LOG] Logotipo adicionado com sucesso")
            logotipo_adicionado = True
        except Exception as e:
            print(f"[AVISO] Arquivo logotipo.png não é uma imagem válida: {str(e)}")
            print(f"[INFO] Substituir img/logotipo.png por uma imagem PNG/JPG real")
    
    if not logotipo_adicionado:
        print(f"[INFO] Cabeçalho da capa criado sem logotipo")
        # Opcional: adicionar texto de placeholder
        # run_header = p_header.add_run("🏥 BANCO DE QUESTÕES MÉDICAS")
        # run_header.bold = True
    
    # Título da capa
    document.add_paragraph("")  # Espaço no topo
    document.add_paragraph("")
    document.add_paragraph("")
    
    capa_title = document.add_paragraph()
    capa_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = capa_title.add_run(f"Banco de Questões - Consulta SQL Específica")
    run.bold = True
    run.font.size = Pt(24)
    
    document.add_paragraph("")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitle.add_run(f"({len(questoes_com_topico)} Questões)")
    run_sub.font.size = Pt(18)
    
    # === SEÇÃO 2: SUMÁRIO ===
    print("[LOG] Criando seção do sumário...")
    document.add_section(WD_SECTION.NEW_PAGE)
    
    # Configurar cabeçalho da seção sumário (sem logotipo)
    section_sumario = document.sections[-1]
    section_sumario.header.is_linked_to_previous = False
    header_sumario = section_sumario.header
    for p in header_sumario.paragraphs:
        p.clear()
    
    # Título do sumário
    sumario_title = document.add_heading("Sumário", level=1)
    sumario_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document.add_paragraph("")
    toc_paragraph = document.add_paragraph()
    add_toc(toc_paragraph)
    
    # === SEÇÃO 3: CONTEÚDO DAS QUESTÕES ===
    print("[LOG] Criando seção do conteúdo das questões...")
    document.add_section(WD_SECTION.NEW_PAGE)
    
    # Adicionar questões organizadas hierarquicamente
    questao_num = 1
    questoes_adicionadas = set() if not permitir_repeticao else None
    
    # Processar cada árvore de tópicos
    for idx_tree, tree in enumerate(topic_trees, 1):
        print(f"[LOG] Processando árvore {idx_tree}: {tree['nome']}")
        
        # Usar função recursiva para adicionar seções hierárquicas
        questao_num = add_topic_sections_recursive(
            document,
            tree,
            reorganized_questions,
            level=1,
            numbering=[idx_tree],
            parent_names=[],
            questao_num=questao_num,
            breadcrumb_raiz=None,  # Não usar breadcrumb_raiz, usar lógica específica
            permitir_repeticao=permitir_repeticao,
            questoes_adicionadas=questoes_adicionadas,
            total_questoes_banco=total_questoes
        )
    
    # Adicionar rodapé
    add_footer_with_text_and_page_number(document)
    
    # Salvar documento
    data_atual = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_filename = f"banco_questoes_sql_{len(questoes_com_topico)}_{data_atual}.docx"
    
    document.save(output_filename)
    print(f"[LOG] Arquivo {output_filename} gerado com sucesso.")
    print(f"[LOG] Total de questões no banco: {len(questoes_com_topico)}")
    
    return output_filename

def gerar_banco_area_especifica(conn, id_topico, total_questoes=1000, permitir_repeticao=True):
    """
    Gera um banco de questões de um tópico específico (qualquer nível na hierarquia).
    
    Args:
        conn: Conexão com o banco de dados
        id_topico: ID do tópico que define a área específica (qualquer nível)
        total_questoes: Número total de questões desejadas
        permitir_repeticao: Se permite questões repetidas
    """
    print(f"[LOG] Gerando banco de questões para tópico específico - Tópico: {id_topico}, {total_questoes} questões...")
    
    cursor = conn.cursor(dictionary=True)
    
    # Primeiro, verificar se o tópico existe e obter seu nome
    cursor.execute("SELECT id, nome FROM topico WHERE id = %s", (id_topico,))
    topico_info = cursor.fetchone()
    
    if not topico_info:
        print(f"[ERRO] Tópico com ID {id_topico} não encontrado!")
        return None
    
    nome_topico = topico_info['nome']
    nome_titulo_limpo = limpar_nome_para_titulo(nome_topico)
    print(f"[LOG] Tópico selecionado: {nome_topico}")
    
    # Informar comportamento de seções baseado no número de questões
    if total_questoes <= 500:
        print(f"[LOG] Banco COMPACTO ({total_questoes} questões): quebras de página apenas para tópicos de NÍVEL 1")
    else:
        print(f"[LOG] Banco EXPANDIDO ({total_questoes} questões): quebras de página para tópicos de NÍVEIS 1, 2 e 3")
    
    # Buscar questões diretamente associadas ao tópico especificado
    # Incluir questões do tópico e de todos os seus descendentes
    print(f"[LOG] Buscando questões associadas ao tópico {id_topico} e seus descendentes...")
    
    # Primeiro, obter todos os descendentes do tópico (incluindo ele próprio)
    cursor.execute("""
        WITH RECURSIVE topico_descendentes AS (
            SELECT id, nome, 1 as nivel
            FROM topico 
            WHERE id = %s
            
            UNION ALL
            
            SELECT t.id, t.nome, td.nivel + 1
            FROM topico t
            INNER JOIN topico_descendentes td ON t.id_pai = td.id
            WHERE td.nivel < 10
        )
        SELECT id FROM topico_descendentes
    """, (id_topico,))
    
    descendentes = cursor.fetchall()
    ids_descendentes = [d['id'] for d in descendentes]
    
    print(f"[LOG] Tópico {id_topico} tem {len(ids_descendentes)} descendentes (incluindo ele próprio)")
    
    if not ids_descendentes:
        print(f"[ERRO] Não foi possível obter descendentes do tópico {id_topico}")
        return None
    
    # Buscar questões associadas a qualquer um dos tópicos descendentes
    format_strings = ','.join(['%s'] * len(ids_descendentes))
    
    query_questoes = f"""
    SELECT DISTINCT
        q.*,
        ROW_NUMBER() OVER (
            ORDER BY SHA2(CONCAT(q.questao_id, 'SEMENTE_FIXA'), 256)
        ) AS ordem
    FROM questaoresidencia q
    INNER JOIN classificacao_questao cq ON q.questao_id = cq.id_questao
    WHERE cq.id_topico IN ({format_strings})
      AND CHAR_LENGTH(q.comentario) >= 500 
      AND q.ano >= 2018
    ORDER BY ordem
    LIMIT %s
    """
    
    print(f"[LOG] Executando consulta SQL para buscar questões do tópico {id_topico}...")
    cursor.execute(query_questoes, tuple(ids_descendentes + [total_questoes]))
    questoes_selecionadas = cursor.fetchall()
    
    print(f"[LOG] Total de questões selecionadas: {len(questoes_selecionadas)}")
    
    if len(questoes_selecionadas) == 0:
        print(f"[ERRO] Nenhuma questão encontrada para o tópico {id_topico}")
        return None
    
    
    # Mapear questões aos tópicos mais específicos possíveis
    print("[LOG] Mapeando questões aos tópicos mais específicos...")
    
    questao_ids = [q['questao_id'] for q in questoes_selecionadas]
    
    # Buscar classificações específicas das questões selecionadas
    format_strings_questoes = ','.join(['%s'] * len(questao_ids))
    format_strings_topicos = ','.join(['%s'] * len(ids_descendentes))
    
    query_topicos_especificos = f"""
    SELECT DISTINCT cq.id_topico, cq.id_questao
    FROM classificacao_questao cq
    WHERE cq.id_questao IN ({format_strings_questoes})
      AND cq.id_topico IN ({format_strings_topicos})
    ORDER BY cq.id_questao, cq.id_topico
    """
    
    cursor.execute(query_topicos_especificos, tuple(questao_ids + ids_descendentes))
    classificacoes_especificas = cursor.fetchall()
    
    print(f"[LOG] Classificações específicas encontradas: {len(classificacoes_especificas)}")
    
    # Criar mapeamento de questão -> tópicos específicos
    questao_topicos_especificos = {}
    for classificacao in classificacoes_especificas:
        questao_id = classificacao['id_questao']
        topico_id = classificacao['id_topico']
        if questao_id not in questao_topicos_especificos:
            questao_topicos_especificos[questao_id] = []
        questao_topicos_especificos[questao_id].append(topico_id)
    
    # Associar cada questão ao tópico mais específico disponível
    questoes_sem_topico = 0
    for q in questoes_selecionadas:
        topicos_especificos = questao_topicos_especificos.get(q['questao_id'], [])
        if topicos_especificos:
            # Usar o primeiro tópico específico encontrado
            q['id_topico'] = topicos_especificos[0]
        else:
            # Fallback: usar o tópico raiz especificado
            q['id_topico'] = id_topico
            questoes_sem_topico += 1
    
    if questoes_sem_topico == 0:
        print(f"[LOG] Todas as questões mapeadas para tópicos específicos")
    else:
        print(f"[LOG] {questoes_sem_topico} questões mapeadas para o tópico raiz (fallback)")
    
    questoes_com_topico = questoes_selecionadas
    print(f"[LOG] Questões com tópico associado: {len(questoes_com_topico)}")
    
    # Verificar se obtivemos o número esperado de questões
    if len(questoes_com_topico) < total_questoes:
        diferenca = total_questoes - len(questoes_com_topico)
        print(f"[AVISO] Obtidas apenas {len(questoes_com_topico)} questões de {total_questoes} solicitadas.")
        print(f"[AVISO] Diferença: {diferenca} questões. Isso pode indicar que não há questões suficientes")
        print(f"[AVISO] no tópico que atendam aos critérios (comentário ≥500 chars, ano ≥2018, etc.)")
    
    # Mostrar distribuição final por área (informativo)
    distribuicao_final = {}
    for q in questoes_com_topico:
        area = q['area']
        distribuicao_final[area] = distribuicao_final.get(area, 0) + 1
    
    print("[LOG] Distribuição final por área:")
    for area, count in distribuicao_final.items():
        print(f"  - {area}: {count} questões")
    
    # Mostrar status final
    if len(questoes_com_topico) == total_questoes:
        print(f"✅ [SUCESSO] Exatamente {total_questoes} questões obtidas!")
    else:
        print(f"⚠️ [AVISO] Obtidas {len(questoes_com_topico)} questões de {total_questoes} solicitadas")
    
    # Organizar questões por tópico
    questions_by_topic = {}
    for q in questoes_com_topico:
        tid = q['id_topico']
        if tid not in questions_by_topic:
            questions_by_topic[tid] = []
        questions_by_topic[tid].append(q)
    
    print(f"[LOG] Questões organizadas por {len(questions_by_topic)} tópicos")
    
    # Construir hierarquia completa dos tópicos utilizados
    print("[LOG] Construindo hierarquia completa dos tópicos...")
    
    topicos_utilizados = list(set([q['id_topico'] for q in questoes_com_topico]))
    topicos_completos = set(topicos_utilizados)
    
    # Para cada tópico utilizado, buscar todos os ancestrais
    for topico_id in topicos_utilizados:
        cursor.execute("""
            WITH RECURSIVE topico_ancestrais AS (
                SELECT id, id_pai, nome, 1 as nivel
                FROM topico 
                WHERE id = %s
                
                UNION ALL
                
                SELECT t.id, t.id_pai, t.nome, ta.nivel + 1
                FROM topico t
                INNER JOIN topico_ancestrais ta ON ta.id_pai = t.id
                WHERE ta.nivel < 10
            )
            SELECT id FROM topico_ancestrais
        """, (topico_id,))
        
        ancestrais = cursor.fetchall()
        for ancestral in ancestrais:
            topicos_completos.add(ancestral['id'])
    
    print(f"[LOG] Tópicos completos (incluindo ancestrais): {len(topicos_completos)}")
    
    # Buscar informações completas dos tópicos
    topicos_completos_list = list(topicos_completos)
    format_strings = ','.join(['%s'] * len(topicos_completos_list))
    
    cursor.execute(f"""
        SELECT id, nome, id_pai
        FROM topico 
        WHERE id IN ({format_strings})
        ORDER BY id
    """, tuple(topicos_completos_list))
    
    topicos_info = {t['id']: t for t in cursor.fetchall()}
    
    # Construir árvore hierárquica a partir do tópico especificado
    def build_topic_tree(topico_id, nivel_atual=1, max_nivel=4):
        if topico_id not in topicos_info:
            return None
            
        topico = topicos_info[topico_id]
        tree_node = {
            'id': topico_id,
            'nome': topico['nome'],
            'nivel': nivel_atual,
            'children': []
        }
        
        if nivel_atual >= max_nivel:
            return tree_node
        
        # Encontrar filhos diretos
        filhos = [t_id for t_id, t_info in topicos_info.items() 
                 if t_info['id_pai'] == topico_id and t_id in topicos_completos]
        
        for filho_id in sorted(filhos):
            child_tree = build_topic_tree(filho_id, nivel_atual + 1, max_nivel)
            if child_tree:
                tree_node['children'].append(child_tree)
        
        return tree_node
    
    # Construir árvore a partir do tópico especificado
    topic_tree = build_topic_tree(id_topico)
    
    if not topic_tree:
        print(f"[ERRO] Não foi possível construir hierarquia para o tópico {id_topico}")
        return None
    
    print(f"[LOG] Árvore hierárquica construída a partir do tópico: {topic_tree['nome']}")
    
    # Reorganizar questões para tópicos de nível 4
    def get_all_descendants(topico_id):
        """Retorna todos os descendentes de um tópico (incluindo ele próprio)"""
        descendants = {topico_id}
        
        filhos = [t_id for t_id, t_info in topicos_info.items() 
                 if t_info['id_pai'] == topico_id]
        
        for filho_id in filhos:
            descendants.update(get_all_descendants(filho_id))
        
        return descendants
    
    def reorganize_questions_for_level4(tree_node, questions_by_topic, reorganized_questions):
        """Reorganiza questões para que tópicos de nível 4 incluam questões de todos os descendentes"""
        
        if tree_node['nivel'] == 4:
            # Este é um tópico de nível 4, coletar questões de todos os descendentes
            all_descendants = get_all_descendants(tree_node['id'])
            todas_questoes = []
            questoes_ids_unicos = set()  # Para evitar duplicatas
            
            for desc_id in all_descendants:
                if desc_id in questions_by_topic:
                    for questao in questions_by_topic[desc_id]:
                        # Verificar se a questão já foi adicionada (evitar duplicatas)
                        if questao['questao_id'] not in questoes_ids_unicos:
                            todas_questoes.append(questao)
                            questoes_ids_unicos.add(questao['questao_id'])
            
            if todas_questoes:
                reorganized_questions[tree_node['id']] = todas_questoes
                print(f"[LOG] Tópico nível 4 '{tree_node['nome']}': {len(todas_questoes)} questões reagrupadas (duplicatas removidas)")
            
        elif tree_node['nivel'] < 4:
            # Para níveis menores que 4, manter questões diretas e processar filhos
            if tree_node['id'] in questions_by_topic:
                reorganized_questions[tree_node['id']] = questions_by_topic[tree_node['id']]
            
            # Processar filhos recursivamente
            for child in tree_node['children']:
                reorganize_questions_for_level4(child, questions_by_topic, reorganized_questions)
    
    # Aplicar reorganização
    reorganized_questions = {}
    reorganize_questions_for_level4(topic_tree, questions_by_topic, reorganized_questions)
    
    print(f"[LOG] Questões reorganizadas para {len(reorganized_questions)} tópicos")
    
    # Criar documento
    document = Document()
    
    # Configurar metadados do documento
    configurar_metadados_documento(document, len(questoes_com_topico), nome_titulo_limpo)
    
    # Configurar estilo padrão
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(3)
    paragraph_format.space_before = Pt(0)
    paragraph_format.line_spacing = 1
    
    # === SEÇÃO 1: CAPA ===
    print("[LOG] Criando seção da capa...")
    
    # Configurar cabeçalho da capa com logotipo
    section_capa = document.sections[0]
    section_capa.header.is_linked_to_previous = False
    header_capa = section_capa.header
    for p in header_capa.paragraphs:
        p.clear()
    
    # Adicionar logotipo no cabeçalho (se disponível)
    img_path = os.path.join(os.path.dirname(__file__), 'img', 'logotipo.png')
    p_header = header_capa.paragraphs[0]
    p_header.clear()
    p_header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    logotipo_adicionado = False
    if os.path.exists(img_path):
        print(f"[LOG] Verificando logotipo: {img_path}")
        run_header = p_header.add_run()
        try:
            Image.open(img_path).verify()
            run_header.add_picture(img_path, width=Inches(3))
            print(f"[LOG] Logotipo adicionado com sucesso")
            logotipo_adicionado = True
        except Exception as e:
            print(f"[AVISO] Arquivo logotipo.png não é uma imagem válida: {str(e)}")
    
    if not logotipo_adicionado:
        print(f"[INFO] Cabeçalho da capa criado sem logotipo")
    
    # Título da capa
    document.add_paragraph("")
    document.add_paragraph("")
    document.add_paragraph("")
    
    capa_title = document.add_paragraph()
    capa_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = capa_title.add_run(f"Banco de Questões - {nome_topico}")
    run.bold = True
    run.font.size = Pt(24)
    
    document.add_paragraph("")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitle.add_run(f"({len(questoes_com_topico)} Questões)")
    run_sub.font.size = Pt(18)
    
    # === SEÇÃO 2: SUMÁRIO ===
    print("[LOG] Criando seção do sumário...")
    document.add_section(WD_SECTION.NEW_PAGE)
    
    section_sumario = document.sections[-1]
    section_sumario.header.is_linked_to_previous = False
    header_sumario = section_sumario.header
    for p in header_sumario.paragraphs:
        p.clear()
    
    sumario_title = document.add_heading("Sumário", level=1)
    sumario_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document.add_paragraph("")
    toc_paragraph = document.add_paragraph()
    add_toc(toc_paragraph)
    
    # === SEÇÃO 3: CONTEÚDO DAS QUESTÕES ===
    print("[LOG] Criando seção do conteúdo das questões...")
    document.add_section(WD_SECTION.NEW_PAGE)
    
    # Adicionar questões organizadas hierarquicamente
    questao_num = 1
    questoes_adicionadas = set() if not permitir_repeticao else None
    
    # No MODO 2, processar diretamente os filhos do tópico raiz como nível 1
    # para que o sumário não mostre o nome da área como tópico principal
    print(f"[LOG] MODO 2: Processando filhos do tópico raiz '{topic_tree['nome']}' como nível 1")
    
    if topic_tree.get('children'):
        # Processar cada filho do tópico raiz como nível 1
        for idx_child, child_tree in enumerate(topic_tree['children'], 1):
            print(f"[LOG] Processando filho {idx_child}: {child_tree['nome']}")
            
            questao_num = add_topic_sections_recursive(
                document,
                child_tree,
                reorganized_questions,
                level=1,
                numbering=[idx_child],
                parent_names=[],
                questao_num=questao_num,
                breadcrumb_raiz=None,
                permitir_repeticao=permitir_repeticao,
                questoes_adicionadas=questoes_adicionadas,
                total_questoes_banco=len(questoes_com_topico)
            )
    else:
        # Se não há filhos, processar o próprio tópico raiz (fallback)
        print(f"[LOG] Tópico raiz '{topic_tree['nome']}' não possui filhos, processando como único tópico")
        questao_num = add_topic_sections_recursive(
            document,
            topic_tree,
            reorganized_questions,
            level=1,
            numbering=[1],
            parent_names=[],
            questao_num=questao_num,
            breadcrumb_raiz=None,
            permitir_repeticao=permitir_repeticao,
            questoes_adicionadas=questoes_adicionadas,
            total_questoes_banco=len(questoes_com_topico)
        )
    
    # Adicionar rodapé
    add_footer_with_text_and_page_number(document)
    
    # Salvar documento
    data_atual = datetime.now().strftime("%Y%m%d_%H%M%S")
    nome_arquivo_limpo = nome_titulo_limpo.replace(" ", "_")
    output_filename = f"banco_questoes_{nome_arquivo_limpo}_{len(questoes_com_topico)}_{data_atual}.docx"
    
    document.save(output_filename)
    print(f"[LOG] Arquivo {output_filename} gerado com sucesso.")
    print(f"[LOG] Total de questões no banco: {len(questoes_com_topico)}")
    
    return output_filename

def gerar_banco_por_instituicao(conn, instituicao, permitir_repeticao=True):
    """
    Gera um banco de questões baseado na instituição (REVALIDA NACIONAL, ENARE ou outra informada) e ano >= 2016.
    Recupera todas as questões que atendam aos critérios, sem cotas por área.
    
    Args:
        conn: Conexão com o banco de dados
        instituicao: 'REVALIDA NACIONAL' ou 'ENARE'
        permitir_repeticao: Se permite questões repetidas
    """
    print(f"[LOG] Gerando banco de questões para {instituicao}...")
    print(f"[LOG] Filtros: instituição LIKE '%{instituicao}%', ano >= 2016, comentário >= 400 caracteres")
    print(f"[LOG] SEM COTAS POR ÁREA - Recuperando todas as questões que atendam aos critérios")
    
    cursor = conn.cursor(dictionary=True)
    
    # Consulta SQL simplificada - sem cotas por área
    query_questoes = f"""
    SELECT 
        q.*
    FROM questaoresidencia q
    WHERE (CHAR_LENGTH(comentario)>400 or (CHAR_LENGTH(comentarioIA)>400 and gabaritoIA=gabarito))
      AND q.ano >= 2016
      AND q.instituicao LIKE '%{instituicao}%'
    ORDER BY q.ano DESC, q.questao_id
    """
    
    print(f"[LOG] Executando consulta SQL simplificada para selecionar questões de {instituicao}...")
    cursor.execute(query_questoes)
    questoes_selecionadas = cursor.fetchall()
    
    print(f"[LOG] Total de questões selecionadas: {len(questoes_selecionadas)}")
    
    if len(questoes_selecionadas) == 0:
        print(f"[ERRO] Nenhuma questão encontrada para {instituicao} com os critérios especificados")
        return None
    
    # Mostrar distribuição por área das questões selecionadas (apenas informativo)
    distribuicao_selecionadas = {}
    for q in questoes_selecionadas:
        area = q['area']
        distribuicao_selecionadas[area] = distribuicao_selecionadas.get(area, 0) + 1
    
    print(f"[LOG] Distribuição por área das questões selecionadas de {instituicao}:")
    for area, count in distribuicao_selecionadas.items():
        print(f"  - {area}: {count} questões")
    
    # Mapear áreas para tópicos raiz (mesmo mapeamento dos modos 1 e 2)
    area_para_topico_raiz = {
        'Cirurgia': 33,
        'Clínica Médica': 100,
        'Pediatria': 48,
        'Ginecologia': 183,
        'Obstetrícia': 218,
        'Medicina Preventiva': 29
    }
    
    print(f"[LOG] Mapeamento área -> tópico raiz: {area_para_topico_raiz}")
    
    # Associar cada questão ao seu tópico raiz baseado na área
    questoes_sem_topico = 0
    for q in questoes_selecionadas:
        area = q['area']
        topico_raiz = area_para_topico_raiz.get(area)
        if topico_raiz:
            q['id_topico'] = topico_raiz
        else:
            print(f"[ERRO] Área '{area}' não mapeada para tópico raiz")
            q['id_topico'] = None
            questoes_sem_topico += 1
    
    if questoes_sem_topico == 0:
        print(f"[LOG] Todas as questões associadas aos tópicos raiz por área")
    else:
        print(f"[ERRO] {questoes_sem_topico} questões não puderam ser associadas a tópicos")
    
    # Obter questões com classificações mais específicas para melhor organização
    questao_ids = [q['questao_id'] for q in questoes_selecionadas]
    format_strings = ','.join(['%s'] * len(questao_ids))
    
    query_topicos_especificos = f"""
    SELECT DISTINCT cq.id_topico, cq.id_questao
    FROM classificacao_questao cq
    WHERE cq.id_questao IN ({format_strings})
    ORDER BY cq.id_questao, cq.id_topico
    """
    
    cursor.execute(query_topicos_especificos, tuple(questao_ids))
    classificacoes_especificas = cursor.fetchall()
    
    print(f"[LOG] Classificações específicas encontradas: {len(classificacoes_especificas)}")
    
    # Criar mapeamento de questão -> tópicos específicos para melhor organização
    questao_topicos_especificos = {}
    for classificacao in classificacoes_especificas:
        questao_id = classificacao['id_questao']
        topico_id = classificacao['id_topico']
        if questao_id not in questao_topicos_especificos:
            questao_topicos_especificos[questao_id] = []
        questao_topicos_especificos[questao_id].append(topico_id)
    
    # Usar tópico mais específico se disponível, senão manter tópico raiz
    for q in questoes_selecionadas:
        topicos_especificos = questao_topicos_especificos.get(q['questao_id'], [])
        if topicos_especificos:
            # Usar o primeiro tópico específico encontrado para melhor organização
            q['id_topico'] = topicos_especificos[0]
        # Se não houver tópico específico, mantém o tópico raiz já definido
    
    questoes_com_topico = questoes_selecionadas
    print(f"[LOG] Questões com tópico associado: {len(questoes_com_topico)}")
    
    # Mostrar distribuição final por área (apenas informativo)
    distribuicao_final = {}
    for q in questoes_com_topico:
        area = q['area']
        distribuicao_final[area] = distribuicao_final.get(area, 0) + 1
    
    print(f"[LOG] Distribuição final por área ({instituicao}):")
    for area, count in distribuicao_final.items():
        print(f"  - {area}: {count} questões")
    
    print(f"✅ [SUCESSO] {len(questoes_com_topico)} questões obtidas para {instituicao}!")
    
    # Obter todos os tópicos únicos das questões
    topicos_utilizados = list(set([q['id_topico'] for q in questoes_com_topico]))
    print(f"[LOG] Tópicos únicos utilizados: {len(topicos_utilizados)}")
    
    # Organizar questões por tópico
    questions_by_topic = {}
    for q in questoes_com_topico:
        tid = q['id_topico']
        if tid not in questions_by_topic:
            questions_by_topic[tid] = []
        questions_by_topic[tid].append(q)
    
    print(f"[LOG] Questões organizadas por {len(questions_by_topic)} tópicos")
    
    # Construir hierarquia completa dos tópicos utilizados (mesmo processo do modo 1)
    print("[LOG] Construindo hierarquia completa dos tópicos...")
    
    topicos_completos = set(topicos_utilizados)
    
    # Para cada tópico utilizado, buscar todos os ancestrais
    for topico_id in topicos_utilizados:
        cursor.execute("""
            WITH RECURSIVE topico_ancestrais AS (
                SELECT id, id_pai, nome, 1 as nivel
                FROM topico 
                WHERE id = %s
                
                UNION ALL
                
                SELECT t.id, t.id_pai, t.nome, ta.nivel + 1
                FROM topico t
                INNER JOIN topico_ancestrais ta ON ta.id_pai = t.id
                WHERE ta.nivel < 10
            )
            SELECT id FROM topico_ancestrais
        """, (topico_id,))
        
        ancestrais = cursor.fetchall()
        for ancestral in ancestrais:
            topicos_completos.add(ancestral['id'])
    
    print(f"[LOG] Tópicos completos (incluindo ancestrais): {len(topicos_completos)}")
    
    # Buscar informações completas dos tópicos
    topicos_completos_list = list(topicos_completos)
    format_strings = ','.join(['%s'] * len(topicos_completos_list))
    
    cursor.execute(f"""
        SELECT id, nome, id_pai
        FROM topico 
        WHERE id IN ({format_strings})
        ORDER BY id
    """, tuple(topicos_completos_list))
    
    topicos_info = {t['id']: t for t in cursor.fetchall()}
    
    # Construir árvores hierárquicas (mesmo processo do modo 1)
    def build_topic_tree(topico_id, nivel_atual=1, max_nivel=4):
        if topico_id not in topicos_info:
            return None
            
        topico = topicos_info[topico_id]
        tree_node = {
            'id': topico_id,
            'nome': topico['nome'],
            'nivel': nivel_atual,
            'children': []
        }
        
        if nivel_atual >= max_nivel:
            return tree_node
        
        # Encontrar filhos diretos
        filhos = [t_id for t_id, t_info in topicos_info.items() 
                 if t_info['id_pai'] == topico_id and t_id in topicos_completos]
        
        for filho_id in sorted(filhos):
            child_tree = build_topic_tree(filho_id, nivel_atual + 1, max_nivel)
            if child_tree:
                tree_node['children'].append(child_tree)
        
        return tree_node
    
    # Encontrar tópicos raiz
    topicos_raiz = []
    for topico_id in topicos_completos:
        if topico_id not in topicos_info:
            print(f"[AVISO] Tópico ID {topico_id} não encontrado em topicos_info, pulando...")
            continue
        topico = topicos_info[topico_id]
        if topico['id_pai'] is None or topico['id_pai'] not in topicos_completos:
            topicos_raiz.append(topico_id)
    
    print(f"[LOG] Tópicos raiz encontrados: {len(topicos_raiz)}")
    
    if len(topicos_raiz) == 0:
        print("[ERRO] Nenhum tópico raiz encontrado. Verificando dados...")
        print(f"[DEBUG] topicos_completos: {len(topicos_completos)}")
        print(f"[DEBUG] topicos_info: {len(topicos_info)}")
        return None
    
    # Construir árvores para cada tópico raiz
    topic_trees = []
    for raiz_id in sorted(topicos_raiz):
        tree = build_topic_tree(raiz_id)
        if tree:
            topic_trees.append(tree)
    
    print(f"[LOG] Árvores construídas: {len(topic_trees)}")
    
    # Definir ordem específica das áreas médicas (mesmo do modo 1)
    ordem_areas = [
        'Cirurgia',
        'Clínica Médica',
        'Pediatria', 
        'Ginecologia',
        'Obstetrícia',
        'Medicina Preventiva'
    ]
    
    # Função para determinar a área de um tópico baseado nas questões
    def get_area_from_topic(tree, questions_by_topic):
        def collect_questions_from_tree(node):
            all_questions = []
            if node['id'] in questions_by_topic:
                all_questions.extend(questions_by_topic[node['id']])
            for child in node.get('children', []):
                all_questions.extend(collect_questions_from_tree(child))
            return all_questions
        
        questoes = collect_questions_from_tree(tree)
        if questoes:
            return questoes[0].get('area', 'Outros')
        return 'Outros'
    
    # Organizar árvores por área
    arvores_por_area = {}
    for tree in topic_trees:
        area = get_area_from_topic(tree, questions_by_topic)
        if area not in arvores_por_area:
            arvores_por_area[area] = []
        arvores_por_area[area].append(tree)
    
    print(f"[LOG] Árvores organizadas por área: {list(arvores_por_area.keys())}")
    
    # Ordenar árvores conforme a sequência desejada
    topic_trees_ordenadas = []
    for i, area in enumerate(ordem_areas, 1):
        if area in arvores_por_area:
            arvores_area = sorted(arvores_por_area[area], key=lambda x: x['nome'])
            topic_trees_ordenadas.extend(arvores_area)
            print(f"[LOG] {i}. Adicionada área '{area}' com {len(arvores_area)} árvore(s)")
    
    # Adicionar áreas não mapeadas no final
    for area, arvores in arvores_por_area.items():
        if area not in ordem_areas:
            arvores_area = sorted(arvores, key=lambda x: x['nome'])
            topic_trees_ordenadas.extend(arvores_area)
            print(f"[LOG] Adicionada área adicional '{area}' com {len(arvores_area)} árvore(s)")
    
    topic_trees = topic_trees_ordenadas
    print(f"[LOG] Árvores reordenadas conforme sequência solicitada: {len(topic_trees)} árvores")
    
    # Reorganizar questões para tópicos de nível 4 (mesmo processo do modo 1)
    def get_all_descendants(topico_id):
        descendants = {topico_id}
        filhos = [t_id for t_id, t_info in topicos_info.items() 
                 if t_info['id_pai'] == topico_id]
        for filho_id in filhos:
            descendants.update(get_all_descendants(filho_id))
        return descendants
    
    def reorganize_questions_for_level4(tree_node, questions_by_topic, reorganized_questions):
        if tree_node['nivel'] == 4:
            all_descendants = get_all_descendants(tree_node['id'])
            todas_questoes = []
            questoes_ids_unicos = set()  # Para evitar duplicatas
            
            for desc_id in all_descendants:
                if desc_id in questions_by_topic:
                    for questao in questions_by_topic[desc_id]:
                        # Verificar se a questão já foi adicionada (evitar duplicatas)
                        if questao['questao_id'] not in questoes_ids_unicos:
                            todas_questoes.append(questao)
                            questoes_ids_unicos.add(questao['questao_id'])
            
            if todas_questoes:
                reorganized_questions[tree_node['id']] = todas_questoes
                print(f"[LOG] Tópico nível 4 '{tree_node['nome']}': {len(todas_questoes)} questões reagrupadas (duplicatas removidas)")
        elif tree_node['nivel'] < 4:
            if tree_node['id'] in questions_by_topic:
                reorganized_questions[tree_node['id']] = questions_by_topic[tree_node['id']]
            for child in tree_node['children']:
                reorganize_questions_for_level4(child, questions_by_topic, reorganized_questions)
    
    # Aplicar reorganização
    reorganized_questions = {}
    for tree in topic_trees:
        reorganize_questions_for_level4(tree, questions_by_topic, reorganized_questions)
    
    print(f"[LOG] Questões reorganizadas para {len(reorganized_questions)} tópicos")
    
    # Criar documento
    document = Document()
    
    nome_titulo_instituicao = limpar_nome_para_titulo(instituicao)
    # Configurar metadados do documento
    configurar_metadados_documento(document, len(questoes_com_topico), nome_titulo_instituicao)
    
    # Configurar estilo padrão
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(3)
    paragraph_format.space_before = Pt(0)
    paragraph_format.line_spacing = 1
    
    # === SEÇÃO 1: CAPA ===
    print("[LOG] Criando seção da capa...")
    
    # Configurar cabeçalho da capa com logotipo
    section_capa = document.sections[0]
    section_capa.header.is_linked_to_previous = False
    header_capa = section_capa.header
    for p in header_capa.paragraphs:
        p.clear()
    
    # Adicionar logotipo no cabeçalho (se disponível)
    img_path = os.path.join(os.path.dirname(__file__), 'img', 'logotipo.png')
    p_header = header_capa.paragraphs[0]
    p_header.clear()
    p_header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    logotipo_adicionado = False
    if os.path.exists(img_path):
        print(f"[LOG] Verificando logotipo: {img_path}")
        run_header = p_header.add_run()
        try:
            Image.open(img_path).verify()
            run_header.add_picture(img_path, width=Inches(3))
            print(f"[LOG] Logotipo adicionado com sucesso")
            logotipo_adicionado = True
        except Exception as e:
            print(f"[AVISO] Arquivo logotipo.png não é uma imagem válida: {str(e)}")
    
    if not logotipo_adicionado:
        print(f"[INFO] Cabeçalho da capa criado sem logotipo")
    
    # Título da capa
    document.add_paragraph("")
    document.add_paragraph("")
    document.add_paragraph("")
    
    capa_title = document.add_paragraph()
    capa_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = capa_title.add_run(f"Banco de Questões - {instituicao}")
    run.bold = True
    run.font.size = Pt(24)
    
    document.add_paragraph("")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitle.add_run(f"({len(questoes_com_topico)} Questões - Ano 2017 em diante)")
    run_sub.font.size = Pt(18)
    
    # === SEÇÃO 2: SUMÁRIO ===
    print("[LOG] Criando seção do sumário...")
    document.add_section(WD_SECTION.NEW_PAGE)
    
    section_sumario = document.sections[-1]
    section_sumario.header.is_linked_to_previous = False
    header_sumario = section_sumario.header
    for p in header_sumario.paragraphs:
        p.clear()
    
    sumario_title = document.add_heading("Sumário", level=1)
    sumario_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document.add_paragraph("")
    toc_paragraph = document.add_paragraph()
    add_toc(toc_paragraph)
    
    # === SEÇÃO 3: CONTEÚDO DAS QUESTÕES ===
    print("[LOG] Criando seção do conteúdo das questões...")
    document.add_section(WD_SECTION.NEW_PAGE)
    
    # Adicionar questões organizadas hierarquicamente
    questao_num = 1
    questoes_adicionadas = set() if not permitir_repeticao else None
    
    # Processar cada árvore de tópicos (mesma estrutura do modo 1)
    for idx_tree, tree in enumerate(topic_trees, 1):
        print(f"[LOG] Processando árvore {idx_tree}: {tree['nome']}")
        
        questao_num = add_topic_sections_recursive(
            document,
            tree,
            reorganized_questions,
            level=1,
            numbering=[idx_tree],
            parent_names=[],
            questao_num=questao_num,
            breadcrumb_raiz=None,
            permitir_repeticao=permitir_repeticao,
            questoes_adicionadas=questoes_adicionadas,
            total_questoes_banco=len(questoes_com_topico)
        )
    
    # Adicionar rodapé
    add_footer_with_text_and_page_number(document)
    
    # Salvar documento
    data_atual = datetime.now().strftime("%Y%m%d_%H%M%S")
    nome_arquivo_limpo = nome_titulo_instituicao.replace(" ", "_")
    output_filename = f"banco_questoes_{nome_arquivo_limpo}_{len(questoes_com_topico)}_{data_atual}.docx"
    
    document.save(output_filename)
    print(f"[LOG] Arquivo {output_filename} gerado com sucesso.")
    print(f"[LOG] Total de questões no banco: {len(questoes_com_topico)}")
    
    return output_filename

if __name__ == "__main__":
    print("=== GERADOR DE BANCO DE QUESTÕES MÉDICAS ===")
    print()
    print("Escolha o modo de geração:")
    print("1 - Banco completo com 6 áreas médicas (Modo original)")
    print("2 - Banco de tópico específico (qualquer nível na hierarquia)")
    print("3 - Banco por instituição (REVALIDA NACIONAL/ENARE/Outra) - Ano 2016 em diante")
    print("4 - Processar questões com comentários incompletos (DeepSeek AI)")
    print("5 - Processar questões específicas por ID (DeepSeek AI)")
    print("6 - Classificar questões sem tópico (DeepSeek AI)")
    print()
    
    # Escolher modo de operação
    try:
        modo = int(input("Digite sua opção (1, 2, 3, 4, 5 ou 6): "))
        if modo not in [1, 2, 3, 4, 5, 6]:
            print("Erro: Opção inválida! Digite 1, 2, 3, 4, 5 ou 6.")
            exit(1)
    except ValueError:
        print("Erro: Digite um número válido (1, 2, 3, 4, 5 ou 6)!")
        exit(1)
    
    # Solicitar número total de questões
    if modo not in [3, 4, 5, 6]:
        try:
            N = int(input("Número total de questões do banco (ex: 1000, 2000, 3000): "))
            if N <= 0:
                print("Erro: N deve ser um número positivo!")
                exit(1)
        except ValueError:
            print("Erro: N deve ser um número inteiro!")
            exit(1)
    
    # Configurar permitir repetição (fixo como False para evitar questões duplicadas)
    permitir_repeticao = False
    
    # Conectar ao banco
    conn = get_connection()
    print("[LOG] Conexão com o banco estabelecida.")
    
    if modo == 1:
        # MODO 1: Banco com 6 áreas médicas (modo original)
        print(f"\n[LOG] MODO 1: Gerando banco com 6 áreas médicas")
        print(f"[LOG] Distribuição proporcional para {N} questões:")
        print(f"  1. Cirurgia: {round(N * 0.2)} questões (20%)")
        print(f"  2. Clínica Médica: {round(N * 0.2)} questões (20%)")
        print(f"  3. Pediatria: {round(N * 0.2)} questões (20%)")
        print(f"  4. Ginecologia: {round(N * 0.1)} questões (10%)")
        print(f"  5. Obstetrícia: {round(N * 0.1)} questões (10%)")
        print(f"  6. Medicina Preventiva: {round(N * 0.2)} questões (20%)")
        print()
        
        gerar_banco_estratificacao_deterministica(conn, N, permitir_repeticao=permitir_repeticao)
        
    elif modo == 2:
        # MODO 2: Banco de tópico específico (qualquer nível)
        print(f"\n[LOG] MODO 2: Gerando banco de tópico específico")
        print()
        print("Códigos dos tópicos raiz das principais áreas:")
        print("  33  - Cirurgia")
        print("  100 - Clínica Médica")
        print("  48  - Pediatria")
        print("  183 - Ginecologia")
        print("  218 - Obstetrícia")
        print("  29  - Medicina Preventiva")
        print()
        print("Ou informe o código de qualquer tópico (raiz ou sub-tópico) desejado.")
        print("O sistema irá buscar questões associadas ao tópico e todos os seus descendentes.")
        print()
        
        try:
            id_topico = int(input("Digite o código do tópico: "))
            if id_topico <= 0:
                print("Erro: O código do tópico deve ser um número positivo!")
                exit(1)
        except ValueError:
            print("Erro: Digite um código válido (número inteiro)!")
            exit(1)
        
        print(f"[LOG] Tópico selecionado: {id_topico}")
        print(f"[LOG] Gerando {N} questões do tópico e seus descendentes...")
        print()
        
        resultado = gerar_banco_area_especifica(conn, id_topico, N, permitir_repeticao=permitir_repeticao)
        
        if not resultado:
            print("[ERRO] Falha na geração do banco de questões!")
            conn.close()
            exit(1)
    
    elif modo == 3:
        # MODO 3: Banco por instituição (REVALIDA/ENARE)
        print(f"\n[LOG] MODO 3: Gerando banco por instituição")
        print()
        print("Instituições disponíveis:")
        print("1:  REVALIDA NACIONAL - Revalidação de diplomas médicos obtidos no exterior")
        print("2:  ENARE - Exame Nacional de Revalidação de Diplomas Médicos Expedidos por Instituições de Educação Superior Estrangeiras")
        print("3:  Outro - Informar o nome da instituição desejada")
        print()
        
        while True:
            try:
                opcao_instituicao = int(input("Digite o número da instituição (1, 2 ou 3): "))
                if opcao_instituicao in [1, 2, 3]:
                    break
                print("Erro: Opção inválida! Digite 1, 2 ou 3.")
            except ValueError:
                print("Erro: Digite um número válido (1, 2 ou 3)!")
        
        if opcao_instituicao == 1:
            instituicao_input = 'REVALIDA NACIONAL'
        elif opcao_instituicao == 2:
            instituicao_input = 'ENARE'
        else:
            instituicao_input = ""
            while not instituicao_input:
                instituicao_input = input("Digite o nome da instituição desejada: ").strip()
                if not instituicao_input:
                    print("Erro: O nome da instituição não pode ser vazio!")
        
        print(f"[LOG] Instituição selecionada: {instituicao_input}")
        print(f"[LOG] Filtros aplicados: ano >= 2016, comentário >= 400 caracteres")
        print(f"[LOG] SEM COTAS POR ÁREA - Recuperando todas as questões que atendam aos critérios")
        print()
        
        resultado = gerar_banco_por_instituicao(conn, instituicao_input, permitir_repeticao=permitir_repeticao)
        
        if not resultado:
            print("[ERRO] Falha na geração do banco de questões!")
            conn.close()
            exit(1)
    
    elif modo == 4:
        # MODO 4: Processar questões com comentários incompletos
        print(f"\n[LOG] MODO 4: Processando questões com comentários incompletos")
        print(f"[LOG] Usando API DeepSeek para análise e justificativa")
        print()

        instituicao_input = ""
        while not instituicao_input:
            instituicao_input = input("Informe o nome (ou parte) da instituição das questões a processar: ").strip()
            if not instituicao_input:
                print("Erro: O nome da instituição não pode ser vazio!")

        # Solicitar RESTO (0 a 4) para permitir processamento paralelo
        try:
            resto = int(input("Informe o RESTO (0-4) para filtrar por questao_id % 5 = RESTO: "))
            if resto not in [0, 1, 2, 3, 4]:
                print("Erro: RESTO deve ser um número entre 0 e 4!")
                conn.close()
                exit(1)
        except ValueError:
            print("Erro: RESTO deve ser um número inteiro entre 0 e 4!")
            conn.close()
            exit(1)

        print(f"[LOG] Instituição: {instituicao_input}")
        print(f"[LOG] Filtrando questões: questao_id % 5 = {resto}")
        processar_questoes_incompletas(conn, instituicao_input, resto)
    
    elif modo == 5:
        # MODO 5: Processar questões específicas por ID e/ou filtros
        print(f"\n[LOG] MODO 5: Processando questões por ID e/ou filtros")
        print(f"[LOG] Usando API DeepSeek para análise e justificativa")
        print()
        
        # Solicitar IDs das questões (opcional)
        questao_ids = None
        ids_input = input("Informe um ou mais IDs de questões (separados por vírgula, ou Enter para não filtrar por ID): ").strip()
        if ids_input:
            try:
                ids_str = [id_str.strip() for id_str in ids_input.split(',')]
                questao_ids = []
                
                for id_str in ids_str:
                    try:
                        questao_id = int(id_str)
                        if questao_id <= 0:
                            print(f"[AVISO] ID inválido ignorado: {id_str} (deve ser positivo)")
                            continue
                        questao_ids.append(questao_id)
                    except ValueError:
                        print(f"[AVISO] ID inválido ignorado: {id_str} (deve ser um número)")
                
                if questao_ids:
                    print(f"[LOG] IDs de questões a processar: {questao_ids}")
                else:
                    questao_ids = None
                    print("[AVISO] Nenhum ID válido fornecido, continuando apenas com filtros")
            except Exception as e:
                print(f"[AVISO] Erro ao processar IDs: {str(e)}, continuando apenas com filtros")
                questao_ids = None
        
        # Solicitar limite (opcional)
        limite = None
        limite_input = input("Informe o número máximo de questões a processar (padrão todas, digite 0 para todas): ").strip()
        if limite_input:
            try:
                limite_val = int(limite_input)
                if limite_val < 0:
                    print("Erro: o número deve ser maior ou igual a zero!")
                    conn.close()
                    exit(1)
                limite = None if limite_val == 0 else limite_val
            except ValueError:
                print("Erro: informe um número inteiro válido!")
                conn.close()
                exit(1)
        
        # Solicitar filtro de instituição (opcional)
        filtro_instituicao = None
        instituicao_input = input("Deseja filtrar por instituição? (pressione Enter para todas): ").strip()
        if instituicao_input:
            filtro_instituicao = instituicao_input
        
        # Solicitar resto módulo 5 (opcional)
        resto_mod5 = None
        resto_input = input("Aplicar filtro questao_id % 5 = RESTO? (Enter para não filtrar): ").strip()
        if resto_input:
            try:
                resto_val = int(resto_input)
                if resto_val not in [0, 1, 2, 3, 4]:
                    print("Erro: RESTO deve ser 0, 1, 2, 3 ou 4!")
                    conn.close()
                    exit(1)
                resto_mod5 = resto_val
            except ValueError:
                print("Erro: RESTO deve ser um número inteiro entre 0 e 4!")
                conn.close()
                exit(1)
        
        # Solicitar filtro de ano mínimo (opcional)
        filtro_ano = None
        ano_input = input("Deseja filtrar por ano mínimo da prova? (Ex: 2018, Enter para todos): ").strip()
        if ano_input:
            try:
                ano_val = int(ano_input)
                filtro_ano = ano_val
            except ValueError:
                print("Erro: ano deve ser um número inteiro válido!")
                conn.close()
                exit(1)
        
        # Validar se pelo menos um critério foi fornecido
        if not questao_ids and limite is None and filtro_instituicao is None and resto_mod5 is None and filtro_ano is None:
            print("[ERRO] Nenhum critério de busca fornecido! Forneça IDs ou pelo menos um filtro.")
            conn.close()
            exit(1)
        
        # Exibir resumo dos filtros
        if questao_ids:
            print(f"[LOG] IDs de questões: {questao_ids}")
        if limite is not None:
            print(f"[LOG] Limite: {limite}")
        if filtro_instituicao:
            print(f"[LOG] Filtro de instituição: {filtro_instituicao}")
        if resto_mod5 is not None:
            print(f"[LOG] Filtro questao_id % 5 = {resto_mod5}")
        if filtro_ano is not None:
            print(f"[LOG] Filtro ano mínimo: {filtro_ano}")
        
        try:
            processar_questoes_por_id(conn, questao_ids=questao_ids, limite=limite, 
                                     filtro_instituicao=filtro_instituicao, 
                                     resto_mod5=resto_mod5, filtro_ano=filtro_ano)
        except Exception as e:
            print(f"[ERRO] Erro ao processar questões: {str(e)}")
            conn.close()
            exit(1)
    
    elif modo == 6:
        # MODO 6: Classificar questões sem tópico
        print(f"\n[LOG] MODO 6: Classificando questões sem tópico")
        print(f"[LOG] Usando API DeepSeek para navegação hierárquica de tópicos")
        print()

        limite = 20
        limite_input = input("Informe o número máximo de questões a classificar (padrão 20, digite 0 para todas): ").strip()
        if limite_input:
            try:
                limite_val = int(limite_input)
                if limite_val < 0:
                    print("Erro: o número deve ser maior ou igual a zero!")
                    conn.close()
                    exit(1)
                limite = None if limite_val == 0 else limite_val
            except ValueError:
                print("Erro: informe um número inteiro válido!")
                conn.close()
                exit(1)
        else:
            limite = 20

        filtro_instituicao = input("Deseja filtrar por instituição? (pressione Enter para todas): ").strip()
        if not filtro_instituicao:
            filtro_instituicao = None

        resto_mod5 = None
        resto_input = input("Aplicar filtro questao_id % 5 = RESTO? (Enter para não filtrar): ").strip()
        if resto_input:
            try:
                resto_val = int(resto_input)
                if resto_val not in [0, 1, 2, 3, 4]:
                    print("Erro: RESTO deve ser 0, 1, 2, 3 ou 4!")
                    conn.close()
                    exit(1)
                resto_mod5 = resto_val
            except ValueError:
                print("Erro: RESTO deve ser um número inteiro entre 0 e 4!")
                conn.close()
                exit(1)

        print(f"[LOG] Limite: {'todas' if limite is None else limite}")
        if filtro_instituicao:
            print(f"[LOG] Filtro de instituição: {filtro_instituicao}")
        if resto_mod5 is not None:
            print(f"[LOG] Filtro questao_id %% 5 = {resto_mod5}")
        filtro_ano = None
        ano_input = input("Deseja filtrar por ano mínimo da prova? (Ex: 2018, Enter para todos): ").strip()
        if ano_input:
            try:
                ano_val = int(ano_input)
                filtro_ano = ano_val
                print(f"[LOG] Filtro ano mínimo: {filtro_ano}")
            except ValueError:
                print("Erro: ano deve ser um número inteiro válido!")
                conn.close()
                exit(1)

        processar_classificacao_questoes_sem_topico(
            conn,
            limite=limite,
            filtro_instituicao=filtro_instituicao,
            resto_mod5=resto_mod5,
            filtro_ano=filtro_ano
        )
    
    conn.close()
    print("\n[LOG] Processo concluído!")
